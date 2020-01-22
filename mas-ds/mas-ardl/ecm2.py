import shutil

import numpy as np
import pandas as pd
import io
import statsmodels.api as sm
import mcModels as mc
import warnings
import copy as cp
import database as databaseConnection
import OneFactorTransformation as oneFactorTransformations
import json
import random
import os

from docx import Document
from docx.shared import Inches
from scipy.stats import shapiro
from statsmodels.graphics.tsaplots import plot_pacf
import matplotlib.pyplot as plt
from statsmodels.stats.outliers_influence import variance_inflation_factor

from statsmodels.tsa.stattools import adfuller

warnings.filterwarnings("ignore", category=FutureWarning)
warnings.filterwarnings("ignore", category=RuntimeWarning)

def wrangle_model_data(file, sheet_name, type='Log-Log'):
    '''
	read the dataset and make date the index of df
	transform data for regression type y~x
	Trim MEV from the front of variable names
	inputs:
    * file - .xlsx file with data
    * sheet_name - name of tab within .xlsx file
    outputs:
    * df - transformed data with Date index
    '''

    transformations = ""
    df = pd.read_excel(file, sheet_name=sheet_name)
    try:
        df.index = df['Date']  # make date column as the index of the dataframe

        df = df.drop('Date', axis=1)
    except:
        print("Your data does not contain Date column")

    df['MEV-BAA-10y Spread'] = df['MEV-BAA CORPORATE BOND YIELD'] - df['MEV-10-YEAR TREASURY NOTE']
    df['MEV-AA-10y Spread'] = df['MEV-AA CORPORATE BOND YIELD'] - df['MEV-10-YEAR TREASURY NOTE']
    df['MEV-B-10y Spread'] = df['MEV-B CORPORATE YIELD'] - df['MEV-10-YEAR TREASURY NOTE']

    if type == 'Level-Log':
        transformations = transformations + " " + type
        df.iloc[:, 2:] = np.log(df.iloc[:, 2:])

    if type == 'Log-Log':
        transformations = transformations + " " + type
        df = np.log(df)

    if type == 'Log-Level':
        transformations = transformations + " " + type
        df.iloc[:, 0] = np.log(df.iloc[:, 0])

    # Rename columns and remove MEV prefix
    cols = []
    for i in df.columns:
        if 'MEV' in i:
            cols.append(i[4:])
        else:
            cols.append(i)

    df.columns = cols

    return [df, transformations]  # returning independent variables

def wrangle_model_data_oneVar(file, sheet_name, date_column="Date"):
    '''
       read the dataset and make date the index of df
       transform data for regression type y~x
       Trim MEV from the front of variable names
       inputs:
       * file - .xlsx file with data
       * sheet_name - name of tab within .xlsx file
       * date_column - The column containing Date which will be used as index
       outputs:
       * regs - dataframe containing all the transformed dataframes for OneFactor
       * df - transformed data with Date index
       '''

    with open('config_old.json') as f:
        configjson = json.load(f)
    df = pd.read_excel(file, sheet_name = sheet_name)

    # df[date_column] = df[date_column].dt.strftime('%Y-%m-%d')
    try:
        df.index = pd.to_datetime(df[date_column], format='%Y-%m-%d') # make date column as the index of the dataframe
        df = df.drop(date_column, axis=1)
    except:
        print("Your data does not contain Date column")

    for i in df.columns:
        if 'MEV-MSA' in i:
            df = df.drop(i, axis=1)

    if configjson['dependentCol']:
        depName = configjson['dependentCol']

    dependentVariableName = str(df.columns[0])

    df = df.dropna(axis=0)
    mev_df = df.iloc[:, 1:]

    diff_df = oneFactorTransformations.difference(df)

    log_df = oneFactorTransformations.log(df)

    log_diff_df = oneFactorTransformations.log_diff(df)

    percent_diff_df = oneFactorTransformations.percent_diff(df)

    regs = pd.concat([mev_df, diff_df, log_df, log_diff_df, percent_diff_df], axis=1)
    regs = regs.replace([np.inf, -np.inf], np.nan)
    regs = regs.drop(regs.index[0])
    regs = regs.dropna(axis=1)

    return [regs, df]


def find_dgp(dep, prin=False):
    '''
    find the data generating process behind a variable
    creating regression for constant,constant trend and random walks models inputs

    input
    * dep - variable to find data generating process of
     * prin - whether to print the full output of the test or not
     output:
     *dparg - string associated with the type of dgp determined, which corresponds to those of the ADF test from statsmodel
    '''
    # Create difference and lag of dependent variable
    delta_dep = dep - dep.shift(1)

    delta_dep.name = 'revenue change'

    delta_dep = delta_dep.dropna(0)

    lag_dep = dep.shift(1)  # shift vertically down

    lag_dep.name = 'revenue lag1'

    # Set up the variables for the random walk regression
    Y = delta_dep.dropna()
    X = lag_dep.dropna()

    # Test significance of the random walk model on the dependent variable
    model = sm.OLS(Y, X)
    results = model.fit()

    sig = results.pvalues[0]

    if prin == True:
        print('Random walk model is significant? ' + str(sig) + ', Pvalue is equal to ' + str(results.pvalues[0]))
    if sig == True:
        dgp = 'Random Walk'
        dgparg = 'nc'

    # Test significance of random walk with drift model on dependent variable
    X = sm.add_constant(X)
    model = sm.OLS(Y, X)
    results = model.fit()

    if (False in results.pvalues < 0.10) == False:
        sig = True
    else:
        sig = False

    if prin == True:
        print('Random walk with drift model is significant ? ' + str(sig) + ' , Pvalues are equal to ' + str(
            results.pvalues[0]) + ' and ' + str(results.pvalues[1]))
    if sig == True:
        dgp = 'Random Walk with drift'
        dgparg = 'c'

    # Test significance of trend with drift model on dependent variable
    time = pd.Series(range(len(Y)))
    time.name = 'time'
    time.index = X.index

    X = pd.concat([X, time], axis=1)
    X = np.asarray(X)  # Convert the input to an array ignoring the index ( one row = one tuple ) in the array

    model = sm.OLS(Y, X)
    results = model.fit()

    # Print results if result is true
    if (False in results.pvalues < 0.10) == False:
        sig = True
    else:
        sig = False

    if prin == True:
        print('Trend with drift model is significant ? ' + str(sig), str(results.pvalues[0]), str(results.pvalues[1]),
              str(results.pvalues[2]))
    if sig == True:
        dgp = 'Trend with drift'
        dgparg = 'ct'
    try:
        if prin == True:
            print(
                'Based on the testing of significance, the ' + dgp + 'model is the most likely data generating process.')
    except:
        if prin == True:
            print('Testing did not indicate significance for any of the lagged models.')
    return dgparg


def test_stationarity(timeseries, prin=False):
    '''

    Test stationarity of time series
    inputs:
    * timeseries - time series of interest
    * prin -  whether to print full results text of find dgp procedure
    outputs:
    * dftest[1] - pvalue of the ADFuller test performed
    '''

    dgp = find_dgp(timeseries, prin)
    dftest = adfuller(timeseries, regression=dgp, maxlag=4)
    dfoutput = pd.Series(dftest[0:4], index=['Test Statistics', 'P-value', '#Lags Used',
                                             'Number of observations used'])  # todo : store the dfoutput?
    for key, value in dftest[4].items():
        dfoutput['Critical Values (%s)' % key] = value
    return dftest[1]


def integration_order(timeseries, alpha=0.05):
    '''

    find the order of integration of a time series
    inputs:
    * timeseries - time series of interest i.e. independent variables
    * alpha = significance level to perform ADF stationarity test at
    outputs :
    * order - order of integration of time series
    '''

    order = -1
    p = 1
    while p > alpha:
        p = test_stationarity(timeseries)
        order += 1
        timeseries = timeseries - timeseries.shift(1)
        timeseries = timeseries.dropna()

    return order

def create_dummy(dep, exclusions):
    '''

    create a dummy variable for regression
    inputs :
    * dep - dependent variable for the regression
    * exclusions - period to be dummied in the variable (i.e. the periods with value 1)
    outputs :
    * base - dependent variable with dummy variables in a dataframe
    '''
    vec = []
    base = dep
    for i in dep.index:
        if i not in exclusions:
            vec.append(0)
        else:
            vec.append(1)
    vec = pd.Series(vec)
    vec.index = dep.index

    base = pd.concat([base, vec], axis=1)
    return base


def exclusions_to_dummies(dep, exclusions):
    '''
    create a matrix of dependent variables and any necessary dummy variables
    inputs:
    * dep - dependent variable
    * exclusions - a list of lists, each list is a list of quarters to be dummied for that dummy variable
    outputs:
    * base - matrix of dependent and dummies
    '''
    basecols = ['Dependent']
    dummyNames = list()
    for i in range(len(exclusions)):
        if i == 0:
            name = 'Dum1'
            dummyNames.append(name)
            basecols.append(name)
            base = create_dummy(dep, exclusions[
                i])  # the number of dummy variables will be equal to the number of exclusions
        else:
            base = create_dummy(base, exclusions[i])
            name = 'Dum' + str(i + 1)
            dummyNames.append(name)
            basecols.append(name)
    base.columns = basecols
    return [base, dummyNames]


def int_filter(df, dep_order, alpha=0.05):
    '''
    filter regression variables for only those with same order of integration as the dependent
    inputs:
    * df - data frame with dependent variable and possible regressors
    * dep_order - order of integration of the dependent variable
    outputs:
    * regs - data frame with variables that satisfied the filter
    '''

    regs = df.iloc[:, 1:]

    for i in regs.columns:
        order = integration_order(regs[i], alpha=alpha)
        if order != dep_order:
            regs = regs.drop(columns=i)

    return regs


def create_basic_design(base, regs, i):
    '''
    create the design matrix for y = XB regression, with no dummies or lags
    used for cointegration test
    :param base: dependent variables and dummies
    :param regs: df with independent variables to be used
    :param i: series of the independent variables to use in the regs df
    outputs:
    X - design matrix for regression
    :return:
    '''
    X = base
    X = X.drop('Dependent', axis=1)
    X[i] = regs.loc[:, i]  # appending one column (ith) independent variable column to the dependent variable column
    return X


def create_onevar_design(base, regs, i):
    '''
    create the design matrix for y = a + XB regression, including dummies
    inputs:
    * base: dependent variable and dummies NOTE: this function creates the lag of dependent and removes first period
    * regs: df with independent variables to be used
    * i: series name of independent variables to use in the regs df
    outputs:
    * X - the design matrix for the regression
    '''
    dep = base.iloc[:, 0]
    X = base
    X = X.drop('Dependent', axis=1)
    X[i] = regs.loc[:, i]
    return (dep, X)


def regress_data(dependent, independent, intercept=True):
    '''
    Perform ordinary least regression on the data
    :param dependent: pandas data series
    :param independent: pandas data frame
    :param intercept: defaults to true but set as false if no intercept is desired
    Output:
    Results have the following important attributes and methods
    * .summary() - shows a summary of the regression results
    * .params - is a pandas series of the parameters fitted by the model
    '''

    if intercept == True:
        X = sm.add_constant(independent)
    results = sm.OLS(dependent, X).fit()
    return results


def residual_stationarity(timeseries):
    '''
    test stationarity of the residuals from model using the non-constant, non-trend version of the Augmented Dickey-Fuller test
    :param timeseries: residuals to be tested
    outputs :
    * dftest[1] - p-value for ADF test
    '''
    dftest = adfuller(timeseries, regression='nc')
    dfoutput = pd.Series(dftest[0:4], index=['Test Statistic', 'p-value', '#Lags Used',
                                             'Number of Observations Used'])  # todo : store the dfoutput?
    for key, value in dftest[4].items():
        dfoutput['Critical Values (%s)' % key] = value
    return dftest[1]


def residual_integration_order(timeseries, alpha=0.05):
    '''
    Find the order of a set of residuals
    Note this uses the constant version of the ADF test

    :param timeseries: The timeseries of interest
    :param alpha: ignificance level to perform ADF Stationarity test at
    outputs:
    order - order of integration of the timeseries
    '''
    order = -1
    p = 1
    while p > alpha:
        p = residual_stationarity(timeseries)
        order += 1
        timeseries = timeseries - timeseries.shift(1)
        timeseries = timeseries.dropna()
    return order


def create_design(base, regs, i):
    '''
    create the design matrix for y = a + XB regression, including y_lag, x, and x_lag as well as dummies
    inputs:
    * base: dependent variable and dummies NOTE: this function creates the lag of dependent and removes first period
    * regs: df with independent variables to be used
    * i: series name of independent variables to use in the regs df
    outputs:
    * X - the design matrix for the regression
    '''
    dep = base.iloc[:, 0]
    base['Dependent_lag'] = base['Dependent'].shift(1)
    base = base.dropna()
    regs_lag = regs.shift(1)
    cols = []
    for j in regs_lag.columns:
        cols.append(j + '_lag')
    regs_lag.columns = cols
    regs = pd.concat([regs, regs_lag], axis=1)
    regs = regs.dropna()
    X = base
    X = X.drop('Dependent', axis=1)
    X[i] = regs.loc[:, i]
    X[i + '_lag'] = regs.loc[:, i + '_lag']
    return (dep, X)


def model_image_save(model, name):
    '''
    save model results to .txt file as [MODEL NAME].txt
    :param model: model returned from sm.OLS()
    :param name: name of the model for the file created
    outputs:
    * save file in pwd
    '''
    str1 = model.summary().as_text()
    name = name + '.txt'
    file1 = open(name, "a")
    file1.write(str1)
    file1.close()

def durbin_watson(model, lags):
    '''
    test model for autocorrelation using Durbin Watson test - this produces n lags test statistics
    inputs:
    * model - model returned from sm.OLS()
    * lags - number of lags to create test statistic
    outputs:
    * dw - vector of test statistics of each lag, in order
    '''
    dw = []
    for i in range(1, lags + 1):
        top = (model.resid - model.resid.shift(i)) ** 2
        top = top.sum()
        bottom = model.resid[i - 1:] ** 2
        bottom = bottom.sum()
        dw.append(top / bottom)
    return dw


def statistical_testing_onevar(base, regs, var):
    '''
         perform statical testing for One Factor
       * base: dependent variable and dummies
       * regs: df with independent variables to be used
       * var: the name of the independent variable
       * adf_alpha: aplha value for ADF Stationarity test
         outputs:
       * params_df: Dataframe containing Coefficient, P-Value, Standard Error, Newey-West p, Newey-West SE, VIF
       * scalar_diagnostics_df: Dataframe containing RSquared, Adjusted, fpval, RMSE, ABS Err, MAE, MAPE,  bg_pval, bp_pval, whit_pval, sw_pval, AIC and DurbinWatson values
       * err_df: Dataframe containing Actual, Predicted, Error, Absolute Error and APE
       * pacf_plot_df: Dataframe containing byte array for PACF plot and plot name
         '''
    dep, X = create_onevar_design(cp.deepcopy(base), regs, var)
    model = regress_data(dep, X, intercept=True)
    vifs = []
    for i in range(X.shape[1]):
        vifs.append(variance_inflation_factor(np.array(X), i))
    vifs = pd.Series(vifs, index=X.columns)
    rsq = pd.Series(model.rsquared, name='R Sq')
    adj_rsq = pd.Series(model.rsquared_adj, name='Adj R Sq')
    fpval = pd.Series(model.f_pvalue, name='F-Test p')
    param_pvals = model.pvalues
    residuals = model.resid
    bse = model.bse
    params = model.params
    resid = model.resid

    rmse = pd.Series(np.sqrt((model.resid ** 2).sum() / len(model.resid)), name='RMSE')
    abs_err = abs(model.resid)
    mae = pd.Series(np.mean(abs_err), name='MAE')
    ape = abs_err / dep
    mape = pd.Series(np.mean(ape), name='MAPE')

    plot_pacf(model.resid, lags=20)
    buf = io.BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)
    imgByteArr = buf.getvalue()
    name = str(var) + " PACF.png"

    PACFPlotName = pd.Series(name, name='PlotName')
    PACFplot = pd.Series(imgByteArr, name='Plot')

    # Serial correlation test
    try:
        bg = sm.stats.diagnostic.acorr_breusch_godfrey(model, nlags=4, store=False)  # Null: no autocorrelation
        bg_pval = pd.Series(bg[1], name='Breusch Godfrey p')
    except:
        bg_pval = pd.Series(np.nan)
    # Heteroscedasticity Test
    try:
        bp = sm.stats.diagnostic.het_breuschpagan(model.resid, model.model.exog)
        bp_pval = pd.Series(bp[1], name='Breusch Pagan p')
    except:
        bp_pval = np.nan
    #         whit = sm.stats.diagnostic.het_white(model.resid, model.model.exog, retres=False) # Null: homoscedastic
    try:
        whit = sm.stats.diagnostic.het_white(model.resid, model.model.exog)
        whit_pval = pd.Series(whit[1], name='White p')
    except:
        whit_pval = pd.Series(np.nan)
    # normality test
    try:
        sw = shapiro(model.resid)  # Null: Residuals are normally distributed
        sw_pval = pd.Series(sw[1], name='Shapiro Wilk p')
    except:
        sw_pval = np.nan

    # Residual stationarity test
    try:
        adf = adfuller(model.resid, regression='nc')
        adf_pval = pd.Series(adf[1], name='ADF p')
    except:
        adf_pval = pd.Series(np.nan, name='ADF p')

    # AIC goodness of fit
    aic = pd.Series(model.aic, name='AIC')

    # Newey-West Standard Errors For Heteroskedasticity Robustness
    new = model.get_robustcov_results(cov_type='HAC', maxlags=1)
    nw_pvals = pd.Series(new.pvalues)
    nwse = pd.Series(new.bse)
    nwse.index = bse.index
    nw_pvals.index = bse.index

    pred = pd.Series(model.predict(), name='Predicted')
    pred.index = dep.index

    err_df = pd.concat([dep, pred, model.resid, abs_err, ape], axis=1)
    err_df.columns = ['Actual', 'Predicted', 'Error', 'Absolute Error', 'APE']

    params_df = pd.concat([params, param_pvals, bse, nw_pvals, nwse, vifs], axis=1)
    params_df.columns = ['Coefficient', 'P-Value', 'Standard Error', 'Newey-West p', 'Newey-West SE', 'VIF']

    scalar_diagnostics_df = pd.concat(
        [rsq, adj_rsq, fpval, rmse, mae, mape, bg_pval, bp_pval, whit_pval, sw_pval, adf_pval, aic], axis=1)

    pacf_plot_df = pd.concat([PACFplot, PACFPlotName], axis=1)
    return params_df, scalar_diagnostics_df, err_df, pacf_plot_df

def statistical_testing(base, regs, var, adf_alpha):
    '''
      perform statical testing for ARDL
    * base: dependent variable and dummies
    * regs: df with independent variables to be used
    * var: the name of the independent variable
    * adf_alpha: aplha value for ADF Stationarity test
      outputs:
    * params_df: Dataframe containing Coefficient, P-Value, Standard Error, Newey-West p, Newey-West SE, VIF
    * scalar_diagnostics_df: Dataframe containing RSquared, Adjusted, fpval, RMSE, ABS Err, MAE, MAPE,  bg_pval, bp_pval, whit_pval, sw_pval, AIC and DurbinWatson values
    * err_df: Dataframe containing Actual, Predicted, Error, Absolute Error and APE
    * pacf_plot_df: Dataframe containing byte array for PACF plot and plot name
      '''
    dep = base['Dependent']
    X = create_basic_design(cp.deepcopy(base), regs, var)
    dep = dep[X.index]
    model = regress_data(dep, X, intercept=True)
    dep_order = residual_integration_order(model.resid, alpha=adf_alpha)
    # creating regression results and diagnostics for each variable
    dep = base['Dependent']
    dep, X = create_design(cp.deepcopy(base), regs, var)

    dep = dep[X.index]
    model = regress_data(dep, X, intercept=True)

    vifs = []
    for i in range(X.shape[1]):
        vifs.append(variance_inflation_factor(np.array(X), i))

    vifs = pd.Series(vifs, index=X.columns)

    # Save Durbin-Watson pvalues up to 4 lags
    dw = durbin_watson(model, 4)
    dw1 = pd.Series(dw[0], name='DurbinWatson1')
    dw2 = pd.Series(dw[1], name='DurbinWatson2')
    dw3 = pd.Series(dw[2], name='DurbinWatson3')
    dw4 = pd.Series(dw[3], name='DurbinWatson4')

    # Here’s R-squared and Adjusted
    rsq = pd.Series(model.rsquared, name='R Sq')
    adj_rsq = pd.Series(model.rsquared_adj, name='Adj R Sq')

    fpval = pd.Series(model.f_pvalue, name='F-Test p')
    param_pvals = model.pvalues
    residuals = model.resid
    bse = model.bse
    params = model.params

    # Calculating error metrics RMSE, ABS Err, MAE, MAPE
    rmse = pd.Series(np.sqrt((model.resid ** 2).sum() / len(model.resid)), name='RMSE')
    abs_err = abs(model.resid)
    mae = pd.Series(np.mean(abs_err), name='MAE')
    ape = abs_err / dep
    mape = pd.Series(np.mean(ape), name='MAPE')

    # converting the PACF plot png to byte array
    plot_pacf(model.resid, lags=20)
    buf = io.BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)
    imgByteArr = buf.getvalue()
    name = str(var) + " PACF.png"

    # saving the byte array with the plot name to the data-frame
    PACFPlotName = pd.Series(name, name='PlotName')
    PACFplot = pd.Series(imgByteArr, name='Plot')

    # Updated stastical tests for both heteroskedasticity tests and exception handling in case they fail
    # Serial correlation test
    try:
        bg = sm.stats.diagnostic.acorr_breusch_godfrey(model, nlags=4, store=False)  # Null: no autocorrelation
        bg_pval = pd.Series(bg[1], name='Breusch Godfrey p')
    except:
        bg_pval = pd.Series(np.nan, name='Breusch Godfrey p')
    # Heteroscedasticity Test
    try:
        bp = sm.stats.diagnostic.het_breuschpagan(model.resid, model.model.exog)
        bp_pval = pd.Series(bp[1], name='Breusch Pagan p')
    except:
        bp_pval = pd.Series(np.nan, name='Breusch Pagan p')
    #         whit = sm.stats.diagnostic.het_white(model.resid, model.model.exog, retres=False) # Null: homoscedastic
    try:
        whit = sm.stats.diagnostic.het_white(model.resid, model.model.exog)
        whit_pval = pd.Series(whit[1], name='White p')
    except:
        whit_pval = pd.Series(np.nan, name='White p')
    # normality test
    try:
        sw = shapiro(model.resid)  # Null: Residuals are normally distributed
        sw_pval = pd.Series(sw[1], name='Shapiro Wilk p')
    except:
        sw_pval = pd.Series(np.nan, name='Shapiro Wilk p')

    # AIC goodness of fit
    aic = pd.Series(model.aic, name='AIC')

    # Newey-West Standard Errors For Heteroskedasticity Robustness
    new = model.get_robustcov_results(cov_type='HAC', maxlags=1)
    nw_pvals = pd.Series(new.pvalues)
    nwse = pd.Series(new.bse)
    nwse.index = bse.index
    nw_pvals.index = bse.index

    pred = pd.Series(model.predict(), name='Predicted')
    pred.index = dep.index

    err_df = pd.concat([dep, pred, residuals, abs_err, ape], axis=1)
    err_df.columns = ['Actual', 'Predicted', 'Error', 'Absolute Error', 'APE']

    params_df = pd.concat([params, param_pvals, bse, nw_pvals, nwse, vifs], axis=1)
    params_df.columns = ['Coefficient', 'P-Value', 'Standard Error', 'Newey-West p', 'Newey-West SE', 'VIF']

    scalar_diagnostics_df = pd.concat(
        [rsq, adj_rsq, fpval, rmse, mae, mape, bg_pval, bp_pval, whit_pval, sw_pval, aic, dw1, dw2, dw3, dw4], axis=1)

    pacf_plot_df = pd.concat([PACFplot, PACFPlotName], axis=1)
    return params_df, scalar_diagnostics_df, err_df, pacf_plot_df


def statistical_diagnostic_filter(stats_Test_Series, diagnostics, bg_alpha=np.nan, bp_alpha=np.nan, whit_alpha=np.nan, sw_alpha=np.nan,
                                  adf_alpha=np.nan):
    '''
        Filtering the indepndent variables/ models that do not pass the statistical tests.
       inputs:
       * stats_Test_Series - tests to be performed
       * diagnostics - dataframe containing result for all the tests performed in statistical_testing method
       * bg_alpha - threshold value for Breusch Godfrey Autocorrelation Test
       * bp_alpha - threshold value for Breusch Pagan Heteroskedasticity Test
       * whit_alpha - threshold value for White Heteroskedasticity Test
       * sw_alpha - threshold value for Shapiro Wilk Residual Normality Test
       * adf_alpha - threshold value for Augmented Dickey Fuller Residual Stationarity Test
       outputs:
       * result - Boolean True if the independent variable/model passes all the tests
       * reasons - If the model fails, the data-frame contains all the reasons of failure.
       '''

    reason = []
    # reason = 'Passes All Tests Specified'
    result = True

    # Breusch Godfrey Test
    if 'Breusch Godfrey Autocorrelation Test' in stats_Test_Series:
        if bg_alpha == None:
            result = result
        else:
            if bg_alpha > diagnostics['Breusch Godfrey p'].iloc[0]:
                result = False
                reason.append('Failed the Breusch Godfrey Autocorrelation Test')
            else:
                result = result
    # Breusch Pagan Test
    if 'Breusch Pagan Heteroskedasticity Test' in stats_Test_Series:
        if bp_alpha == None:
            result = result
        else:
            if bp_alpha > diagnostics['Breusch Pagan p'].iloc[0]:
                result = False
                reason.append('Failed the Breusch Pagan Heteroskedasticity Test')
            else:
                result = result
    # White Test
    if 'White Heteroskedasticity Test' in stats_Test_Series:
        if whit_alpha == None:
            result = result
        else:
            if 'White p' in diagnostics:
                if whit_alpha > diagnostics['White p'].iloc[0]:
                    result = False
                    reason.append('Failed the White Heteroskedasticity Test')
                else:
                    result = result

    # Shapiro Wilk Test
    if 'Shapiro Wilk Residual Normality Test' in stats_Test_Series:
        if sw_alpha == None:
            result = result
        else:
            if sw_alpha > diagnostics['Shapiro Wilk p'].iloc[0]:
                result = False
                reason.append('Failed the Shapiro Wilk Residual Normality Test')
            else:
                result = result

    # Augmented Dickey Fuller Test
    if 'Augmented Dickey Fuller Residual Stationarity Test' in stats_Test_Series:
        if adf_alpha == None:
            result = result
        else:
            if 'ADF p' in diagnostics.columns:
                if adf_alpha < diagnostics['ADF p'].iloc[0]:
                    result = False
                    reason.append('Failed the Augmented Dickey Fuller Residual Stationarity Test')
                else:
                    result = result

    if result:
        reason.append('Passes All Tests Specified')
    return result, reason

def one_var_param_filter(param_df, param_alpha):
    '''
        Filtering the indepndent variables/ models that do not pass Parameter Significance Test.
         inputs:
         * param_df - dataframe containing result Pvalue for a independent variable obtained after sm.OLS()
         * param_alpha - threshold value for Parameter Significance Test.
         outputs:
         * result - Boolean True if the independent variable/model passes all the tests
         * reasons - If the model fails, the data-frame contains all the reasons of failure.
    '''
    reason = []
    result = True
    if param_alpha == None:
        result = result
    else:
        for i in range(len(param_df.index)):
            if 'const' and 'Dum' not in param_df.index[i]:
                if param_df['P-Value'].iloc[i] > param_alpha:
                    result = False
                    reason.append('Failed the Parameter Significance Test')
                    break
                else:
                    result = result
            else:
                result = result

    if result:
        reason.append('Passes All Tests Specified')

    return result, reason

def run_stats_tests(dummies, transformations, stats_Test_Series, connections, runId, base, regs, typ, param_alpha, bg_alpha=np.nan,
                    bp_alpha=np.nan, whit_alpha=np.nan, sw_alpha=np.nan,
                    adf_alpha=np.nan):
    '''
         Perform regression on each independent variable, perform statistical test and store the result to the database
         * dummies - list of names of the Dummy Variables
         * transformations - the transformation applied to the dependent variable.
         * stats_Test_Series - statistical tests to be performed
         * connection - connection object to perform querying and updating the database
         * runId - runId of the job run.
         * base - dependent variable and dummies
         * regs - df with independent variables to be used
         * typ - the type of model
         * param_alpha - threshold value for Parameter Significance Test.
         * bg_alpha - threshold value for Breusch Godfrey Autocorrelation Test
         * bp_alpha - threshold value for Breusch Pagan Heteroskedasticity Test
         * whit_alpha - threshold value for White Heteroskedasticity Test
         * sw_alpha - threshold value for Shapiro Wilk Residual Normality Test
         * adf_alpha - threshold value for Augmented Dickey Fuller Residual Stationarity Test
         outputs:
         * reasons - If the model fails, the data-frame contains all the reasons of failure.
         * candidates - Candidates / independent variables which passed the test.
    '''
    modelOutputID = 0
    result1 = False
    candidates = []
    reasons = []
    for i in regs.columns:
        reasonList = []
        if typ == 'OneVar':
            modelId, independentVariableId, modelOutputID = databaseConnection.dbCreateModelIdIndependentVarAndOutputId(
                runId=runId, name=str(i), transformations=transformations, Connection=connections)
            params_df, scalar_diagnostics_df, err_df, pacf_plot_df = statistical_testing_onevar(base, regs, i)
            databaseConnection.model_result_save_in_db(scalar_diagnostic=scalar_diagnostics_df, paramsdf=params_df,
                                                       pacf_plot=pacf_plot_df ,name=i, connection=connections, runId=runId, modelId=modelId,
                                                       modelOutputID=modelOutputID,
                                                       independentVariableId=independentVariableId, dummies=dummies)
            result1, reason1 = statistical_diagnostic_filter(stats_Test_Series, scalar_diagnostics_df, bg_alpha=bg_alpha, bp_alpha=bp_alpha,
                                                           whit_alpha=whit_alpha, sw_alpha=sw_alpha,
                                                           adf_alpha=adf_alpha)
            if result1 == True:
                if 'Parameter Significance Test' in stats_Test_Series:
                    result1, reason2 = one_var_param_filter(params_df, param_alpha)
                    if result1 == True:
                        candidates.append(i)
                        reasonList.append(reason1)
                    else:
                        reasonList.append(reason2)
            else:
                reasonList.append(reason1)
                result3, reason3 = one_var_param_filter(params_df, param_alpha)
                if result3 == False:
                    reasonList.append(reason3)
        elif typ == 'ARDL':
            modelId, independentVariableId, modelOutputID = databaseConnection.dbCreateModelIdIndependentVarAndOutputId(
                runId=runId, name=str(i), transformations=transformations, Connection=connections)
            params_df, scalar_diagnostics_df, err_df, pacf_plot_df = statistical_testing(base, regs, i,
                                                                                         adf_alpha)
            databaseConnection.model_result_save_in_db(scalar_diagnostic=scalar_diagnostics_df, paramsdf=params_df,
                                                       pacf_plot=pacf_plot_df, name=i, connection=connections,
                                                       runId=runId, modelId=modelId, modelOutputID=modelOutputID,
                                                       independentVariableId=independentVariableId, dummies=dummies)
            result1, reason1 = statistical_diagnostic_filter(stats_Test_Series, scalar_diagnostics_df, bg_alpha=bg_alpha, bp_alpha=bp_alpha,
                                                           whit_alpha=whit_alpha, sw_alpha=sw_alpha,
                                                           adf_alpha=adf_alpha)
            if result1 == True:
                if 'Parameter Significance Test' in stats_Test_Series:
                    result1, reason2 = one_var_param_filter(params_df, param_alpha)
                    if result1 == True:
                        candidates.append(i)
                        reasonList.append(reason1)
                    else:
                        reasonList.append(reason2)
            else:
                reasonList.append(reason1)
                result3, reason3 = one_var_param_filter(params_df, param_alpha)
                if result3 == False:
                    reasonList.append(reason3)
        databaseConnection.updateModelOutputForResult(connections, modelOutputID, result1, json.dumps(reasonList))
        reasons.append(json.dumps(reasonList))
    reasons = pd.Series(reasons, index=regs.columns)
    return reasons, candidates

def create_backtest(runId, modelId, connection, X, beta, pq1, figname, target, l=9):
    '''
    create backtesting results and saves .png file for log-log model
    inputs:
    * X - design matrix for regression
    * beta - parameters from regression
    * pq1 - first forecast period
    * figname - name of figure for file name
    * target - dependent variable target
    outputs:
    * MAPE - MAPE of dynamic backtest plotted
    * saves plot to pwd as [figname].png
    '''
    plt.clf()
    index = target.index.tolist()
    pqloc = target.index.get_loc(pq1)
    hist_X = X.loc[:pq1]
    post_X = X.loc[pq1:]
    post_X = post_X.iloc[0:l]
    y1 = post_X['Dependent_lag'].iloc[0]
    inds = beta.index
    inds = beta.index.drop('Dependent_lag')
    inds = inds.drop('const')
    y_hat_values = []
    for j in range(l):
        y_hat = beta['const']
        for i in inds:
            y_hat = y_hat + beta[i] * post_X[i].iloc[j]
        if j == 0:
            y_hat = y_hat + beta['Dependent_lag'] * y1
        else:
            y_hat = y_hat + beta['Dependent_lag'] * y_hat_values[-1]
        y_hat_values.append(y_hat)
    y_hat_values_copy = y_hat_values
    y_hat_values = pd.Series(np.exp(y_hat_values))


    if np.inf in y_hat_values.values:
        y_hat_values = np.log(y_hat_values_copy)
        y_hat_values = pd.Series(np.exp(y_hat_values))

    actual = np.exp(target.iloc[pqloc:pqloc + l])

    if np.inf in actual.values:
        target_log = np.log(target.iloc[pqloc:pqloc + l])
        actual = np.exp(target_log)

    y_hat_values.index = actual.index

    error = actual - y_hat_values
    perc_err = error / actual
    ape = abs(perc_err)
    MAPE = np.mean(ape)
    figname = figname + ' MAPE = ' + str(MAPE * 100)[:5] + '%'
    act = pd.Series(np.exp(target.iloc[pqloc - 1]))
    act.index = pd.Series(target.index[pqloc - 1])

    y_hat_values = pd.concat([act, y_hat_values])
    actual = np.exp(target.iloc[0:pqloc + l])

    y_hat_values = y_hat_values.replace([np.inf, -np.inf], np.nan)
    createDictionary(connection, actual, y_hat_values, figname, pq1, modelId)
    return MAPE


def backtesting(runId, connection, candidates, base, regs, dates, long_dates, target="", l=9, ):
    '''
    plot and compile backtesting results for candidate models
    inputs:
    * candidates - list of variables to backtest models for
    * base: dependent variable and dummies
    * regs: df with independent variables to be used
    * dates - list of pq1 periods to start 9Q backtest
    * long_dates - list of pq1 periods to start full length backtest
    * target - target dependent variable
    outputs:
    * mape - matrix of variables and MAPE values for various forecasting periods
    * saves mape matrix to candidate backtesting results.xlsx
    * saves matrix of variables and MAPE values for full backtest to long backtsting results.xlsx file
    '''
    dep = base['Dependent']
    mapes = []
    mape = pd.DataFrame(index=candidates, columns=dates)

    long_mape = pd.DataFrame(index=candidates, columns=long_dates)

    for i in candidates:
        dep, X = create_design(base, regs, i)
        dep = dep[X.index]
        model = regress_data(dep, X, intercept=True)
        beta = model.params
        modelId = connection.getModelId(runId, i)


        for pq1 in dates:
            figname = str(runId) + '_' + str(modelId) + '_9Q Backtest_' + str(pq1)
            mape.loc[i, pq1] = create_backtest(runId, modelId, connection, X, beta, pq1, figname, base['Dependent'])

        for pq1 in long_dates:
            l = len(X.loc[pq1:].index)
            figname = str(runId) + '_' + str(modelId) + '_Full History Backtest_' + str(pq1)
            long_mape.loc[i, pq1] = create_backtest(runId, modelId, connection, X, beta, pq1, figname,
                                                    base['Dependent'], l=l)
    mape.to_excel('Candidate Backtesting MAPE Results.xlsx')
    long_mape.to_excel('Candidate Long Backtesting MAPE Results.xlsx')

    return [mape, long_mape]


def createDictionary(connection, actual, predicted, figname, date, modelId, type='BackTestJson'):
    '''
        create Dictionary for BackTest graph json and save the json to the database
        inputs:
        * connection - the database connection object for updating or querying the database
        * actual - the dataframe containing actual data-points
        * predicted - the dataframe containing predicted data-points
        * figname - name of figure for file name
        * date - the index of the graph
        * modelId - the model/independentVariable id for which the graph is being generated
    '''
    listActualPredicted = []
    listActual = []
    listPredicted = []

    for index, value in actual.iteritems():
        listActual.append({'date': index.isoformat(), 'value': value})

    for index, value in predicted.iteritems():
        listPredicted.append({'date': index.isoformat(), 'value': value})

    actualDict = {'id': 'actual', 'values': listActual}

    predictedDict = {'id': 'predicted', 'values': listPredicted}

    listActualPredicted.extend([actualDict, predictedDict])

    resultDict = {'name': figname, 'Date': date, 'values': listActualPredicted}

    dictjson = json.dumps(resultDict)
    if 'Out of Time' in figname:
        type = 'OutOfTime'

    databaseConnection.saveJson(connection, modelId, figname, dictjson, jsontype=type)


def recursive_reg(dep, X, n, varname):
    '''
    perform recursive regression on model
    inputs:
    * dep - dependent variable
    * X - design matrix
    * n - number of latest periods to use for recursive regression
    * varname - name of the x variable in the model
    outputs:
    * params - parameter values for the recursive regresison trials
    * ps - pvalues for the recursive regression trials
    next steps - change funciton so that a date can be passed in lieu of n
    '''
    # each iteration will generate a set of pvalues and params, then plot each of the pvalues and params.
    ps = pd.DataFrame()
    params = pd.DataFrame()
    for i in range(n, len(dep)):
        dep_trim = dep[0:i]
        X_trim = X.iloc[0:i, :]
        model = mc.regress_data(dep_trim, X_trim, intercept=True)
        params[X_trim.index[-1]] = model.params
        ps[X_trim.index[-1]] = model.pvalues
    return params, ps


def getJsonFromDictionary(keyName, valueName, dictionary, figname, connection, runId, variableName, jsonType):
    '''
           create json from Dictionary for graph and save the json to the database
           inputs:
           * dictionary - dictionary containing the graph data-points
           * figname - name of figure / graph
           * connection - the database connection object for updating or querying the database
           * runId - the model/independentVariable id for which the graph is being generated
           * variableName - name of the independent variable
           * jsonType - type of json to be generated this should be 'BackTestJson/OutOfTime/RegressionParam/RegressionPval/SensitivityTest/StressTestCompare/StressTestForecast'
       '''
    listActual = []
    listResult = []
    for key, value in dictionary.items():
        listActual.append({'date': key.isoformat(), 'value': value})

    actualDict = {'id': 'actual', 'values': listActual}

    listResult.append(actualDict)

    resultDict = {'name': figname, 'values': listResult}

    dictJSON = json.dumps(resultDict)

    modelId = databaseConnection.getModelId(connection, runId, variableName)
    databaseConnection.saveJson(connection, modelId, figname, dictJSON, jsontype=jsonType)


def create_scenario(mev_full):
    '''
    create mev forecasting scenario
    inputs:
    * mev_full - df of mev data with date as first column
    outputs
    * mev_full - df with date index and no MEV_ prefix on varnames
    '''

    mev_full.index = mev_full.iloc[:, 0]
    mev_full = mev_full.drop(mev_full.columns[0], axis=1)
    for i in mev_full.columns:
        if 'MEV' not in i:
            mev_full = mev_full.drop(i, axis=1)
    cols = []
    for i in mev_full.columns:
        cols.append(i[4:])
    mev_full.columns = cols
    return mev_full


def create_forecast(X, beta, pq1, target_spot, l=22):
    '''
    create the forecast for stress test
    inputs:
    * X - design matrix for the model
    * beta - parameters of the model
    * pq1 - first quarter to be forecasted
    * target spot - spot value of the dependent variable
    * l - length of forecast desired
    outputs:
    * y_hat_values - predicted outcome of stress test
    '''
    X = X.iloc[0:l]
    y1 = target_spot
    y_hat = beta[0] + beta[-2] * X.iloc[:, 0] + beta[-1] * X.iloc[:, 1]
    for j in range(l):
        if j == 0:
            y_hat[j] = y_hat[j] + beta[-3] * y1
        else:
            y_hat[j] = y_hat[j] + beta[-3] * y_hat[j - 1]
    y_hat_values = np.exp(y_hat)
    return y_hat_values


def build_scenario(filename, shtm, shtb, shta, shts, beta, pq0, pq1, i, l=22):
    '''
    build forecast for scenarios in the stress test
    inputs:
    * filename - name of the .xlsx file where the data is located
    * shtm - name of the tab with the model development data
    * shtb - name of the tab with the base scenario forecast
    * shta - name of the tab with the adverse scenario forecast
    * shts - name of the tab with the severe scenario forecast
    * beta - parameter vector for the model
    * pq0 - spot period
    * pq1 - first forecasted period
    * i - model variable name
    * l - length of forecast
    outputs:
    * forecast df - data frame with last 4 periods actuals and the base, adverse, and severe scenario forecasts
    '''
    df, _ = wrangle_model_data(filename, sheet_name=shtm)
    full_dep = df.iloc[:, 0]
    full_dep.name = 'Dependent'
    forecast_df = pd.DataFrame()
    base_full = pd.read_excel(filename, sheet_name=shtb)
    b = create_scenario(base_full)
    adverse_full = pd.read_excel(filename, sheet_name=shta)
    a = create_scenario(adverse_full)
    severe_full = pd.read_excel(filename, sheet_name=shts)
    s = create_scenario(severe_full)
    target_spot = full_dep[pq0]
    count = 0
    for scen in [b, a, s]:
        count += 1
        X = pd.DataFrame()
        X[i] = scen[i]
        for j in X.index:
            if j not in scen.index:
                X = X.drop(j)
        X.index = scen.index

        lagname = i + ' 1Q Lag'

        lag = scen[i].shift(1)
        X[lagname] = lag
        X = X.loc[pq1:]
        X = np.log(X)

        result = create_forecast(X, beta, pq1, target_spot, l=l)
        forecast_df[str(count)] = result
    forecast_df.columns = ['Base', 'Adverse', 'Severe']
    forecast_df = pd.concat([np.exp(full_dep), forecast_df], axis=1)
    forecast_df = forecast_df.iloc[-(l + 4):]
    for scenario in ['Base', 'Adverse', 'Severe']:
        forecast_df[scenario].iloc[-(l + 1)] = np.exp(full_dep.iloc[-1])
    return forecast_df


def is_nan(x):
    '''
       check if the value is nan
       inputs:
       * X - the value to check for nan
       outputs:
       * boolean, True if the input is nan else false
       '''
    return (x is np.nan or x != x)

def getJSONForStressTestforecast(name, dictionary):
    '''
        create json from Dictionary for graph and save the json to the database
        inputs:
        * name - name of figure for file name
        * dictionary - dictionary containing values for one of the stress test
        output:
        * actualJson - json containing the the result for single stress test
    '''

    listActual = []

    for key, value in dictionary.items():
        if not is_nan(value):
            listActual.append({'date': key.isoformat(), 'value': value})

    actualJson = {'id': name, 'values': listActual}

    return actualJson


def createJSONForStressTestForecast(dataframe, name):
    '''
        create json from Dictionary for graph and save the json to the database
        inputs:
        * dataframe - dataframe containing the entier result of the stress test
        * name - name of the plot
        output:
        * finalResultJSON - json containing the entier result for stress test
    '''
    outputList = []
    stressTestDict = dataframe.to_dict('dict')

    for key, value in stressTestDict.items():
        intermediateResultDict = getJSONForStressTestforecast(key, value)
        outputList.append(intermediateResultDict)

    finalDict = {'name': name, 'values': outputList}

    finalResultJSON = json.dumps(finalDict)
    return finalResultJSON

def stress_test_plot(filename, shtm, shtb, shta, shts, short_list, pq0, pq1, base, regs, connection, runId, bottom="", top=""):
    '''
    plot stress testing results for candidate models
    inputs
    * filename - name of the .xlsx file where the data is located
    * shtm - name of the tab with the model development data
    * shtb - name of the tab with the base scenario forecast
    * shta - name of the tab with the adverse scenario forecast
    * shts - name of the tab with the severe scenario forecast
    * short_list - list of candidate variables for forecasting
    * pq0 - spot period
    * pq1 - first forecasted period
    * base: dependent variable and dummies
    * regs: df with independent variables to be used
    * bottom - bottom of y-axis of desired forecast graphs
    * top - top of y-axis of desired forecast graphs
    outputs:
    * saves stress test forecast to .png files
    '''

    for i in short_list:
        varname = i
        dep = base['Dependent']
        _, X = create_design(base, regs, i)
        dep = dep[X.index]
        model = mc.regress_data(dep, X, intercept=True)
        beta = model.params

        forecast = build_scenario(filename, shtm, shtb, shta, shts, beta, pq0, pq1, i)
        # print("ForeCast : ", forecast)
        figname = i + ' Stress Test Forecast'
        fig, ax = plt.subplots(ncols=1, nrows=1, figsize=(10, 6))

        if top != "":
            if bottom != "":
                plt.ylim(top=top, bottom=bottom)
            else:
                plt.ylim(top=top)
        else:
            if bottom != "":
                plt.ylim(bottom=bottom)

        plt.setp(ax.get_xticklabels(), rotation=45)
        j = 0
        colors = ['Black', 'Green', 'Blue', 'Red']
        for i in forecast.columns:
            ax.plot(forecast.index, forecast[i], color=colors[j])
            j += 1
        #         ax.legend(loc = 'best')
        ax.set_title(figname)
        vals = ax.get_yticks()
        ax.set_yticklabels(['${0:,.0f}'.format(x) for x in vals])
        json = createJSONForStressTestForecast(forecast, figname)
        modelId = databaseConnection.getModelId(connection, runId, varname)
        databaseConnection.saveJson(connection, modelId, figname, json, jsontype='StressTestForecast')

def wrangle_forecast_data(file, sheet_name):
    '''
    create the mev forecasting data frame
    inputs:
    * file - name of the .xlsx file where the data is located
    * sheet_name - name of the tab with the forecast data scenario
    outputs:
    * df - mev data frame with date index and MEV_ prefix dropped
    '''
    # read the dataset and make date the index of the df
    df = pd.read_excel(file, sheet_name=sheet_name)
    df.index = df['Date']
    df = df.drop('Date', axis=1)
    # Rename columns and remove MEV prefix
    cols = []
    for i in df.columns:
        if 'MEV' in i:
            cols.append(i[4:])
        else:
            cols.append(i)
    df.columns = cols

    return df

def createJSONForStressTestComparison(compare, forecast, figname):
    '''
        create json for StressTestComparison graph
        inputs:
        * compare - compare dataframe
        * forecast - forecast dataframe
        * figname - name of the plot
        output:
        * finalResultJSON - json containing the entier result for stress test
    '''
    finalList = []
    mergedDataFrame = pd.concat([forecast, compare], axis=1, join_axes=[forecast.index])

    dictionary = mergedDataFrame.to_dict()
    # print("Dictionary : ", dictionary)

    for key, value in dictionary.items():
        intermediateResultDict = getJSONForStressTestforecast(key, value)
        finalList.append(intermediateResultDict)

    finalDict = {'name': figname, 'values': finalList}

    finalResultJSON = json.dumps(finalDict)
    return finalResultJSON


def stress_test_compare(filename, shtm, shtb, shta, shts, shtc, short_list, pq0, pq1, base, regs, dep, connection, runId, bottom="",
                        top=""):
    '''
    compare stress test forecasts with those of the current model
    inputs:
    * filename - name of the .xlsx file where the data is located
    * shtm - name of the tab with the model development data
    * shtb - name of the tab with the base scenario forecast
    * shta - name of the tab with the adverse scenario forecast
    * shts - name of the tab with the severe scenario forecast
    * shtc - name of the tab with the comparison forecasts from the current model
    * short_list - list of candidate variables for forecasting
    * pq0 - spot period
    * pq1 - first forecasted period
    * base: dependent variable and dummies
    * regs: df with independent variables to be used
    * dep - dependent variable
    * bottom - bottom of y-axis of desired forecast graphs
    * top - top of y-axis of desired forecast graphs
    outputs:
    * forecast_tbl - forecast metrics for candidate models
    * compare_tbl - forecast metrics for current model
    * saves plots of each forecast to .png file in pwd
    '''

    writer = pd.ExcelWriter('Stress Test Forecast.xlsx', engine='xlsxwriter')
    compare = wrangle_forecast_data(filename, shtc)
    full_df = wrangle_forecast_data(filename, shtm)
    full_dep = full_df.iloc[:, 0]
    full_dep.name = 'Dependent'
    target_spot = full_dep[pq0]
    compare = compare
    forecast_tbl = pd.DataFrame()
    compare_tbl = pd.DataFrame()
    if len(short_list) > 0:
        for i in short_list:
            varname = i
            _, X = create_design(base, regs, i)
            dep = dep[X.index]
            model = mc.regress_data(dep, X, intercept=True)
            beta = model.params
            forecast = build_scenario(filename, shtm, shtb, shta, shts, beta, pq0, pq1, i)

            figname = i + ' Stress Test Comparison'
            fig, ax = plt.subplots(ncols=1, nrows=1, figsize=(10, 6))
            if top != "":
                if bottom != "":
                    plt.ylim(top=top, bottom=bottom)
                else:
                    plt.ylim(top=top)
            else:
                if bottom != "":
                    plt.ylim(bottom=bottom)

            plt.setp(ax.get_xticklabels(), rotation=45)
            j = 0
            colors = ['Black', 'Green', 'Blue', 'Red']
            forecast.columns = ['Actual', 'Model Base', 'Model_Adverse', 'Model_Severe']
            forecast_tbl = pd.DataFrame()
            for k in forecast.columns:
                if j > 0:
                    actual = target_spot
                    pq4 = forecast[k].iloc[4]
                    pq9 = forecast[k].iloc[9]
                    pq20 = forecast[k].iloc[20]
                    cagr9 = (pq9 / actual) ** (4 / 9) - 1
                    cagr20 = (pq20 / actual) ** (4 / 20) - 1

                    actual = '${0:,.0f}'.format(actual)
                    pq4 = '${0:,.0f}'.format(pq4)
                    pq9 = '${0:,.0f}'.format(pq9)
                    pq20 = '${0:,.0f}'.format(pq20)
                    cagr9 = '{:.2%}'.format(cagr9)
                    cagr20 = '{:.2%}'.format(cagr20)
                    forecast_tbl[k] = pd.Series([actual, pq4, pq9, pq20, cagr9, cagr20])
                ax.plot(forecast.index, forecast[k], color=colors[j])
                j += 1
            forecast_tbl = forecast_tbl.transpose()
            forecast_tbl.columns = ['Actual', 'PQ4', 'PQ9', 'PQ20', '9Q CAGR', '20Q CAGR']

            for col in compare.columns:
                compare.loc[full_dep.index[-1], col] = target_spot

            compare = compare.sort_index(axis=0)

            colors = ['Green', 'Blue', 'Red']
            j = 0
            compare_tbl = pd.DataFrame()
            for k in compare.columns:
                ax.plot(compare.index, compare[k], color=colors[j], linestyle='--')

                actual = target_spot
                pq4 = compare[k].iloc[4]
                pq9 = compare[k].iloc[9]
                pq20 = compare[k].iloc[20]
                cagr9 = (pq9 / actual) ** (4 / 9) - 1
                cagr20 = (pq20 / actual) ** (4 / 20) - 1

                actual = '${0:,.0f}'.format(actual)
                pq4 = '${0:,.0f}'.format(pq4)
                pq9 = '${0:,.0f}'.format(pq9)
                pq20 = '${0:,.0f}'.format(pq20)
                cagr9 = '{:.2%}'.format(cagr9)
                cagr20 = '{:.2%}'.format(cagr20)
                compare_tbl[k] = pd.Series([actual, pq4, pq9, pq20, cagr9, cagr20])
                j += 1
            compare_tbl = compare_tbl.transpose()
            compare_tbl.columns = ['Actual', 'PQ4', 'PQ9', 'PQ20', '9Q CAGR', '20Q CAGR']
            #         ax.legend(loc = 'best')
            ax.set_title(figname)
            vals = ax.get_yticks()
            ax.set_yticklabels(['${0:,.0f}'.format(x) for x in vals])
            json = createJSONForStressTestComparison(compare, forecast, figname)
            modelId = databaseConnection.getModelId(connection, runId, varname)
            databaseConnection.saveJson(connection, modelId, figname, json, jsontype='StressTestCompare')
            forecast_tbl.to_excel(writer, sheet_name=i[:25] + " New")
            compare_tbl.to_excel(writer, sheet_name=i[:25] + ' Old')
        writer.save()
    return forecast_tbl, compare_tbl

def create_rands(size, n):
    '''
    generate vector 1xsize random numbers between -n and n
    inputs:
    * size - size of vector desired
    * n - number for random numbers between -n and n
    outputs:
    * rands - series of randomly generated numbers
    '''
    rands = []
    for i in range(size):
        x = random.randrange(-n * 1000, n * 1000, 1) / 1000
        rands.append(x)
    return pd.Series(rands)

def create_random_forecast(X, beta, y1, xspot, xmu, xsig, l=12):
    '''
    create random forecasts based on random variance added to mev forecast
    x_t = x_{t-1} + historical mean(x') +  random decimal (-3,3) * historical standard deviation(x')
    where ' denotes a difference
    inputs:
    * X - design matrix of the model
    * beta - parameters of the model
    * xspot - spot value of the regressor variable x
    * xmu - historical mean of the diff of the x variable
    * xsig - historical standard deviation of the diff of the x variable
    l - length of forecast
    outputs:
    * y_hat_values - model forecast for random scenario
    '''
    X = X.iloc[0:l]
    xsigrands = create_rands(l, 3)
    xlagsigrands = create_rands(l, 3)
    xsigrands.index = X.index
    xlagsigrands.index = X.index
    x = []
    xlag = []
    for i in range(l):
        if i == 0:
            x.append(xspot + xmu + xsigrands[i] * xsig)  # + xmu)
            xlag.append(xspot)
        else:
            x.append(x[-1] + xmu + xsigrands[i] * xsig)  # + xmu)
            xlag.append(x[-1])

    y_hat = beta[0] + beta[-2] * pd.Series(x) + beta[-1] * pd.Series(xlag)
    for j in range(l):
        if j == 0:
            y_hat[j] = y_hat[j] + beta[-3] * y1
        else:
            y_hat[j] = y_hat[j] + beta[-3] * y_hat[j - 1]

    y_hat_values = np.exp(y_hat.astype(float))
    y_hat_values = pd.Series(y_hat_values)
    y_hat_values.index = X.index
    return y_hat_values

def build_random_scenario(filename, shtm, shtb, beta, pq0, pq1, i, l=12):
    '''
    build random forecast for the model
    inputs:
    * filename - name of the .xlsx file where the data is located
    * shtm - name of the tab with the model development data
    * shtb - name of the tab with the base scenario forecast
    * beta - parameters of the model
    * pq0 - spot period
    * pq1 - first forecasted period
    * i - variable for the model
    * l - length of forecast period
    outputs:
    * forecast_df - random model forecast
    '''
    forecast_df = pd.DataFrame()
    base_full = pd.read_excel(filename, sheet_name=shtb)
    scen = create_scenario(base_full)
    count = 0
    X = pd.DataFrame()
    X[i] = scen[i]

    X.index = scen.index
    lagname = i + ' 1Q Lag'
    lag = scen[i].shift(1)
    X[lagname] = lag

    df, _ = wrangle_model_data(filename, shtm)
    dep = df.iloc[:, 0]
    dep.name = 'Dependent'
    y1 = np.log(dep[pq0])
    X = np.log(X)
    diff = X - X.shift(1)
    diff = diff.loc[df.index]
    xsig = np.std(diff[i])
    xmu = np.mean(diff[i])
    xspot = X[i].loc[pq0]
    X = X.loc[pq1:]
    result = create_random_forecast(X, beta, y1, xspot, xmu, xsig)
    forecast_df['Rand'] = result

    forecast_df = pd.concat([dep, forecast_df], axis=1)
    forecast_df = forecast_df.iloc[-(l + 4):]
    forecast_df['Rand'].iloc[-(l + 1)] = forecast_df['Dependent'].iloc[-(l + 1)]
    return forecast_df


def createJSONForSensitivityTest(sensitivityDataframe, figname):
    '''
           create json for SensitivityTest graph
           inputs:
           * sensitivityDataframe - sensitivity dataframe
=           * figname - name of the plot
           output:
           * finalResultJSON - json containing result for sensitivity test
       '''
    finalList = []
    dictionary = sensitivityDataframe.to_dict()

    for key, value in dictionary.items():
        intermediateResultDict = getJSONForStressTestforecast(key, value)
        finalList.append(intermediateResultDict)

    finalDict = {'name': figname, 'values': finalList}

    finalResultJSON = json.dumps(finalDict)
    return finalResultJSON

def create_sensitivity(filename, shtm, shtb, base, regs, short_list, pq0, pq1, connection, runId, l=50):
    '''
    create sensitivity testing for list of candidate models
    inputs:
    * filename - name of the .xlsx file where the data is located
    * shtm - name of the tab with the model development data
    * shtb - name of the tab with the base scenario forecast
    * base: dependent variable and dummies
    * regs: df with independent variables to be used
    * short_list - list of candidate variables for forecasting
    * pq0 - spot period
    * pq1 - first forecasted period
    * l - number of forecasts to randomly generate
    outputs:
    * saves plot of sensitivity analysis to .png in pwd
    '''
    resdf = pd.DataFrame(columns=['Mean', 'Min', 'Max'])
    df,_ = wrangle_model_data(filename, shtm)
    dep = df.iloc[:, 0]
    dep.name = 'Dependent'
    dep = np.log(dep)
    for i in short_list:
        varname = i
        _, X = create_design(base, regs, i)
        X.index = X.index.date
        dep = dep[X.index]
        # added
        dep.reindex(X.index)
        model = regress_data(dep, X, intercept=True)
        beta = model.params
        figname = i + ' Dynamic Sensitivity Testing'
        fig, ax = plt.subplots(ncols=1, nrows=1, figsize=(10, 6))
        ends = []
        sensitivityDataframe = pd.DataFrame()
        for k in range(l):
            j = 0
            forecast = build_random_scenario(filename, shtm, shtb, beta, pq0, pq1, i)
            if k > 0:
                forecastToMerge = forecast.iloc[:, 1:2]
                name = 'Rand'+str(k)
                forecastToMerge.columns = [name]
                forecastToMerge.index = forecast.index
            else:
                forecastToMerge = forecast
            sensitivityDataframe = pd.concat([sensitivityDataframe, forecastToMerge], axis=1, join_axes=[forecast.index])
            ends.append(forecast['Rand'].iloc[-1])
            colors = ['black', 'red']
            for col in forecast.columns:
                ax.plot(forecast.index, forecast[col], color=colors[j])
                j += 1
        try:
            plt.ylim(top=1.25 * np.max(np.max(forecast)), bottom=0.75 * np.min(np.min(forecast)))
        except:
            pass
        plt.setp(ax.get_xticklabels(), rotation=45)
        ax.set_title(figname)
        vals = ax.get_yticks()
        ax.set_yticklabels(['${0:,.0f}'.format(x) for x in vals])
        for x in vals:
            ax.axhline(y=x, color='black', linewidth=0.2)
        json = createJSONForSensitivityTest(sensitivityDataframe, figname)
        modelId = databaseConnection.getModelId(connection, runId, varname)
        databaseConnection.saveJson(connection, modelId, figname, json, jsontype='SensitivityTest')
        resdf.loc[i] = ['${0:,.0f}'.format(np.mean(ends)), '${0:,.0f}'.format(np.min(ends)),
                        '${0:,.0f}'.format(np.max(ends))]
    writer = pd.ExcelWriter('Sensitivity Metrics.xlsx', engine='xlsxwriter')
    resdf.to_excel(writer)
    writer.save()

def oot_backtesting(candidates, base, regs, pq0, dates, full_base, full_regs, target, connection, runId, l=9):
    '''
    perform out of time testing on the list of candidate models
    inputs:
    * candidates - list of models to perform oot testing on
    * base: dependent variable and dummies for intended development period for oot test
    * regs: df with independent variables for intended development period for oot test
    * pq0 - spot period
    * dates - dates to perform backtesting on
    * full base - full dependent variables and dummies
    * full_regs - full independent variables
    * target - model target actuals
    * l - length of forecast for backtesting
    outputs:
    * mape - matrix of MAPE values for forecasts indexed by pq1 period
    '''
    dep = base['Dependent']
    mapes = []
    mape = pd.DataFrame(index=candidates, columns=dates)
    for i in candidates:
        _, full_X = create_design(full_base, full_regs, i)
        _, X = create_design(base, regs, i)
        dep = dep[X.index]
        model = regress_data(dep, X, intercept=True)
        beta = model.params
        for pq1 in dates:
            figname = i + '_' + str(pq0)[:11] + ' Out of Time ' + str(pq1)
            modelId = databaseConnection.getModelId(connection, runId, i)
            mape.loc[i, pq1] = create_backtest(runId, modelId, connection, full_X, beta, pq1, figname, target)
    mape.to_excel('Out of Time MAPE Results.xlsx')
    return mape

def out_of_time(candidates, base, regs, pq0, dates, connection, runId):
    '''
    create out of time testing for list of candidate variables
    inputs:
    * candidates - list of models to perform oot testing on
    * base: dependent variable and dummies
    * regs: df with independent variables
    * pq0 - spot period to use for out of time regressions
    * dates - pq1 periods to perform backtesting on
    outputs:
    * mape - matrix of MAPE values for forecasts indexed by pq1 period
    '''
    full_dep = base['Dependent']
    full_base = base
    full_regs = regs
    base = base[:pq0]
    regs = regs[:pq0]
    mape = oot_backtesting(candidates, base, regs, pq0, dates, full_base, full_regs, full_dep, connection, runId, l=9)
    return mape


def compile_results(short_list):
    '''
    create a word doc of the key results for the final candidate models
    inputs:
    * short_list - list of candidates to compile to document
    outputs:
    * saves .docx file to pwd with results compiled and organized
    '''
    try:
        os.remove('short_list.docx')
    except:
        pass

    document = Document()
    document.add_heading('Document Title', 0)
    for folder in short_list:
        document.add_heading(folder, level=1)
        ls = os.listdir()
        varfiles = []
        for i in ls:
            if (folder in i) and ('VS. ' + folder not in i):
                varfiles.append(i)
        modelfiles = []
        backfiles = []
        outfiles = []
        stressfiles = []
        sensfiles = []
        paramfiles = []
        pfiles = []
        pacffiles = []
        for file in varfiles:
            if file.endswith(".txt"):
                modelfiles.append(file)
            if 'PACF' in file:
                pacffiles.append(file)
            if 'Backtest' in file:
                backfiles.append(file)
            if 'Out of Time' in file:
                outfiles.append(file)
            if 'Stress Test' in file:
                stressfiles.append(file)
            if 'Sensitivity' in file:
                sensfiles.append(file)
            if 'param.png' in file:
                paramfiles.append(file)
            if 'pval.png' in file:
                pfiles.append(file)
        document.add_heading('Model Form', level=2)
        for file in modelfiles:
            file1 = open(file, "r")
            graph = file1.read()
            p = document.add_paragraph(graph)
        for file in pacffiles:
            document.add_picture(file, width=Inches(7))
        document.add_heading('Backtesting', level=2)
        for file in backfiles:
            document.add_picture(file, width=Inches(7))
        document.add_heading('Out of Time Testing', level=2)
        for file in outfiles:
            document.add_picture(file, width=Inches(7))
        document.add_heading('Stress Testing', level=2)
        for file in stressfiles:
            document.add_picture(file, width=Inches(7))
        document.add_heading('Sensitivity Testing', level=2)
        for file in sensfiles:
            document.add_picture(file, width=Inches(7))
        document.add_heading('Recursive', level=2)
        document.add_heading('Parameter Plots', level=3)
        for file in paramfiles:
            document.add_picture(file, width=Inches(7))
        document.add_heading('P-Value Plots', level=3)
        for file in pfiles:
            document.add_picture(file, width=Inches(7))
        document.add_page_break()
    document.save('short_list.docx')

def copy_output(folder):
    '''
    copy docx and xlsx files to folder
    '''
    try:
        os.mkdir(folder)
    except:
        pass
    for i in os.listdir():
        if i.endswith('.xlsx'):
            shutil.copy(i, folder)
        if i.endswith('.docx'):
            shutil.copy(i, folder)

def version_output(folder):
    '''
    organize output into folders
    '''

    os.mkdir(folder)
    backtestloc = folder + '/backtests'
    os.mkdir(backtestloc)
    stloc = folder + '/stress_test'
    os.mkdir(stloc)
    sensloc = folder + '/sensitivity_analysis'
    os.mkdir(sensloc)
    ootloc = folder + '/out_of_time'
    os.mkdir(ootloc)

    for i in os.listdir():
        if 'Out of Time' in i:
            shutil.move(i, ootloc)
        if 'Backtest' in i:
            shutil.move(i, backtestloc)
        if 'Stress Test' in i:
            shutil.move(i, stloc)
        if 'Sensitivity' in i:
            shutil.move(i, sensloc)
            try:
                shutil.move(i, folder)
            except:
                pass
