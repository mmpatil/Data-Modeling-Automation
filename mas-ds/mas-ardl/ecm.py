#!/usr/bin/env python
"""
TODO: Document the module.
Provides classes and functionality for SOME_PURPOSE
"""

#######################################
# Any needed from __future__ imports  #
# Create an "__all__" list to support #
#   "from module import member" use   #
#######################################

__all__ = [
    # Constants
    # Exceptions
    # Functions
    """
        * wrangle_model_data
        * find_dgp
        * test_stationarity
        * integration_order
        * create_dummy
        * exclusions_to_dummies
        * int_filter
        * create_design
        * model_image_save
        * durbin_watson
        * statistical_testing
        * create_backtest
        * backtesting
        * recursive_reg
        * create_scenario
        * create_forecast
        * build_scenario
        * stress_test_plot
        * wrangle_forecast_data
        * stress_test_compare
        * create_rands
        * build_random_scenario
        * create_random_forecast
        * create_sensitivity
        * compile_results
        * copy_output
        * version_output
        * out_of_time
        """
    # ABC "interface" classes
    # ABC abstract classes
    # Concrete classes
]

#######################################
# Module metadata/dunder-names        #
#######################################

__author__ = 'Michael J. McLaughlin'
__copyright__ = 'Copyright 2019, all rights reserved'
__status__ = 'Development'

#######################################
# Standard library imports needed     #
#######################################

# Uncomment this if there are abstract classes or "interfaces" 
#   defined in the module...
# import abc
import subprocess
import os
import sys

#######################################
# Third-party imports needed          #
#######################################
import pandas as pd
import matplotlib

matplotlib.use(
    'Agg')  # Troubleshooting - see here https://stackoverflow.com/questions/41814254/qxcbconnection-error-when-python-matplotlib-run-as-cron-job
import matplotlib.pyplot as plt
import statsmodels.api as sm
from statsmodels.tsa.stattools import adfuller
from scipy.stats import shapiro
import shutil as shutil
import mcModels as mc
import numpy as np
import datetime as dt
from PIL import Image, ImageDraw, ImageFont
from sklearn.linear_model import LassoCV, LassoLarsCV, LassoLarsIC
from statsmodels.graphics.tsaplots import plot_pacf
from docx import Document
from docx.shared import Inches
import random
import copy as cp


#######################################
# Local imports needed                #
#######################################

#######################################
# Initialization needed before member #
#   definition can take place         #
#######################################

#######################################
# Module-level Constants              #
#######################################

#######################################
# Custom Exceptions                   #
#######################################

#######################################
# Module functions                    #
#######################################
def wrangle_model_data(file, sheet_name, type='Level-Level'):
    '''
    read the dataset and make date the index of the df
    transform data for regression type y~x 
    Trim MEV from the front of variable names
    inputs:
    * file - .xlsx file with data
    * sheet_name - name of tab within .xlsx file
    outputs:
    * df - transformed data with Date index
    '''

    df = pd.read_excel(file, sheet_name=sheet_name)
    try:
        df.index = df['Date']
        df = df.drop('Date', axis=1)
    except:
        print('ERROR: Your data does not contain "Date" column')
    df['MEV-BAA-10y Spread'] = df['MEV-BAA CORPORATE BOND YIELD'] - df['MEV-10-YEAR TREASURY NOTE']
    df['MEV-AA-10y Spread'] = df['MEV-AA CORPORATE BOND YIELD'] - df['MEV-10-YEAR TREASURY NOTE']
    df['MEV-B-10y Spread'] = df['MEV-B CORPORATE YIELD'] - df['MEV-10-YEAR TREASURY NOTE']
    if type == 'Level-Log':
        df.iloc[:, 2:] = np.log(df.iloc[:, 2:])

    if type == 'Log-Log':
        # Turn the df into logspace, except for the dummy variables

        df = np.log(df)

    if type == 'Log-Level':
        df.iloc[:, 0] = np.log(df.iloc[:, 0])
    # Rename columns and remove MEV prefix
    cols = []
    for i in df.columns:
        if 'MEV' in i:
            cols.append(i[4:])
        else:
            cols.append(i)

    df.columns = cols

    return df


def find_dgp(dep, prin=False):
    '''
    find the data generating process behind a variable
    creating regressions for constant, constant trend, and random walk models
    inputs:
    * dep - variable to find data generating process of
    * prin - whether to print the full output of the test or not
    output: 
    * dgparg - string associated with type of dgp determined, which corresponds to those of the ADF test from statsmodels
    '''
    # Create difference and lag of dependent variable
    delta_dep = dep - dep.shift(1)
    delta_dep.name = 'revenue change'
    lag_dep = dep.shift(1)
    lag_dep.name = 'revenue lag1'

    # Set up the variables for the random walk regression 
    Y = delta_dep.dropna()
    X = lag_dep.dropna()

    # Test significance of random walk model on the dependent variable
    model = sm.OLS(Y, X)
    results = model.fit()
    sig = results.pvalues[0] <= 0.10
    if prin == True:
        print('Random walk model is significant? ' + str(sig) + ', Pvalue is equal to ' + str(results.pvalues[0]))
    if sig == True:
        dgp = 'Random Walk'
        dgparg = 'nc'
    # Test significance of random walk with drift model on dependent variable
    X = sm.add_constant(X)
    model = sm.OLS(Y, X)
    results = model.fit()
    if (False in results.pvalues < 0.10) == False:
        sig = True
    else:
        sig = False
    if prin == True:
        print('Random walk with drift model is significant? ' + str(sig) + ', Pvalues are equal to ' + str(
            results.pvalues[0]) + ' and ' + str(results.pvalues[1]))
    if sig == True:
        dgp = 'Random walk with drift'
        dgparg = 'c'
    # Test significance of trend with drift model on dependent variable
    time = pd.Series(range(len(Y)))
    time.name = 'time'
    time.index = X.index
    X = pd.concat([X, time], axis=1)
    X = np.asarray(X)

    model = sm.OLS(Y, X)
    results = model.fit()
    # Prin results if prin = True
    if (False in results.pvalues < 0.10) == False:
        sig = True
    else:
        sig = False
    if prin == True:
        print('Trend with drift model is significant? ' + str(sig) + ', Pvalues are equal to ' + str(
            results.pvalues[0]) + ', ' + str(results.pvalues[1]) + ' ,and ' + str(results.pvalues[2]))
    if sig == True:
        dgp = 'Trend with drift'
        dgparg = 'ct'
    try:
        if prin == True:
            print(
                'Based on the testing of significance, the ' + dgp + ' model is the most likely data generating process.')
    except:
        if prin == True:
            print('Testing did not indicate significance for any of the lagged models.')
    return dgparg


# test the stationarity of timeseries
# requires the find_dgp function initialized in earlier cell
def test_stationarity(timeseries, prin=False):
    '''
    test the stationarity of a time series
    inputs:
    * timeseries - time series of interest
    * prin - whether to print full results text of find dgp procedure
    outputs:
    * dftest[1] - pvalue of the ADFuller Test performed
    '''
    dgp = find_dgp(timeseries, prin)
    dftest = adfuller(timeseries, regression=dgp, maxlag=4)
    dfoutput = pd.Series(dftest[0:4], index=['Test Statistic', 'p-value', '#Lags Used', 'Number of Observations Used'])
    for key, value in dftest[4].items():
        dfoutput['Critical Value (%s)' % key] = value
    return dftest[1]


def residual_stationarity(timeseries):
    '''
    test stationarity of residuals from model using the non-constant, non-trend version of the Augmented Dickey-Fuller test
    inputs:
    * timeseries - residuals to be tested
    outputs:
    * dftest[1] - p-value for  ADF test
    '''
    dftest = adfuller(timeseries, regression='nc')
    dfoutput = pd.Series(dftest[0:4], index=['Test Statistic', 'p-value', '#Lags Used', 'Number of Observations Used'])
    for key, value in dftest[4].items():
        dfoutput['Critical Value (%s)' % key] = value
    return dftest[1]


# test the integration order of timeseries
def integration_order(timeseries, alpha=0.05):
    '''
    find the order of integration of a time series
    inputs:
    * timeseries - time series of interest
    * alpha - significance level to perform ADF Stationarity test at
    outputs:
    * order - order of integration of the time series
    '''
    order = -1
    p = 1
    while p > alpha:
        p = test_stationarity(timeseries)
        order += 1
        timeseries = timeseries - timeseries.shift(1)
        timeseries = timeseries.dropna()
    return order


# test the integration order of timeseries
def residual_integration_order(timeseries, alpha=0.05):
    '''
    find the order of integration of a set of residuals
    Note this uses the constant version of the ADF test 
    inputs:
    * timeseries - time series of interest
    * alpha - significance level to perform ADF Stationarity test at
    outputs:
    * order - order of integration of the time series
    '''
    order = -1
    p = 1
    while p > alpha:
        p = residual_stationarity(timeseries)
        order += 1
        timeseries = timeseries - timeseries.shift(1)
        timeseries = timeseries.dropna()
    return order


def create_dummy(dep, exclusions):
    '''
    create a dummy variable for regression 
    inputs:
    * dep - dependent variable of th regression
    * exclusions - the periods to be dummied in the variable (ie. the periods with value 1)
    outputs:
    * base - dependent variable with dummy variables in a data frame 
    '''
    vec = []
    base = dep
    for i in dep.index:
        if i not in exclusions:
            vec.append(0)
        else:
            vec.append(1)
    vec = pd.Series(vec)
    vec.index = dep.index

    base = pd.concat([base, vec], axis=1)
    return base


def exclusions_to_dummies(dep, exclusions):
    '''
    create a matrix of the dependnent variables and any necessary dummy variables
    inputs:
    * dep - dependent variable
    * exclusions - a list of lists, each list is a list of quarters to be dummied for that dummy variable
    outputs:
    * base - matrix of dependent and dummies
    '''
    basecols = ['Dependent']

    for i in range(len(exclusions)):
        if i == 0:
            basecols.append('Dum1')
            base = create_dummy(dep, exclusions[i])
        else:
            base = create_dummy(base, exclusions[i])
            name = 'Dum' + str(i + 1)
            basecols.append(name)
    print(base.columns)
    base.columns = basecols
    print(base)
    return base


def int_filter(df, dep_order, alpha=0.05):
    '''
    filter regression variables for only those with same order of integration as the dependent
    inputs: 
    * df - data frame with dependent variable and possible regressors
    * dep_order - order of integration of the dependent variable
    outputs:
    * regs - data frame with variables that satisfied the filter
    '''
    regs = df.iloc[:, 1:]
    for i in regs.columns:
        order = integration_order(regs[i], alpha=alpha)
        if order != dep_order:
            regs = regs.drop(columns=[i])
    return regs


# create the design matrix based on the
# * base: dependent variable and dummies NOTE: this function creates the lag of dependent and removes first period
# * regs: df with independent variables to be used
# * i: series name of independent variables to use in the regs df
def create_design(base, regs, i):
    '''
    create the design matrix for y = a + XB regression, including y_lag, x, and x_lag as well as dummies
    inputs: 
    * base: dependent variable and dummies NOTE: this function creates the lag of dependent and removes first period
    * regs: df with independent variables to be used
    * i: series name of independent variables to use in the regs df
    outputs:
    * X - the design matrix for the regression
    '''
    dep = base.iloc[:, 0]
    base['Lag_Y'] = base['Dependent'].shift(1)
    base = base.dropna()
    regs_lag = regs.shift(1)
    cols = []
    for j in regs_lag.columns:
        cols.append(j + '_lag')
    regs_lag.columns = cols
    regs = pd.concat([regs, regs_lag], axis=1)
    regs = regs.dropna()
    X = base
    X = X.drop('Dependent', axis=1)
    X[i] = regs.loc[:, i]
    X[i + '_lag'] = regs.loc[:, i + '_lag']
    return (X)


def create_basic_design(base, regs, i):
    '''
    create the design matrix for y = XB regression, with no dummies or lags
    used for the cointegration test
    inputs:
    * base: dependent variable and dummies 
    * regs: df with independent variables to be used
    * i: series name of independent variables to use in the regs df
    outputs:
    X - design matrix for regression 
    '''
    X = base
    X = X.drop('Dependent', axis=1)
    X[i] = regs.loc[:, i]
    return (X)


# write the model results to a .txt file
def model_image_save(model, name):
    '''
    save model results to .txt file as [MODEL NAME].txt
    inputs:
    * model - model returned from sm.OLS()
    * name - name of the model for the file created
    outputs:
    * saves file in pwd
    '''
    str1 = model.summary().as_text()
    # for object name file1. 
    name = name + '.txt'
    file1 = open(name, "a")
    file1.write(str1)
    file1.close()


def durbin_watson(model, lags):
    '''
    test model for autocorrelation using Durbin Watson test - this produces n lags test statistics
    inputs:
    * model - model returned from sm.OLS()
    * lags - number of lags to create test statistic
    outputs:
    * dw - vector of test statistics of each lag, in order
    '''
    dw = []
    for i in range(1, lags + 1):
        top = (model.resid - model.resid.shift(i)) ** 2
        top = top.sum()
        bottom = model.resid[i - 1:] ** 2
        bottom = bottom.sum()
        dw.append(top / bottom)
    return dw


def statistical_testing(base, regs, adf_alpha=0.05, param_alpha=0.10, bg_alpha=0.05, white_alpha=0.05, sw_alpha=0.05):
    '''
    create regressions and run diagnostics
    filter out regressions that do not pass tests
    save statistical testing results to .xlsx file
    tests performed:
    * Breusch Godfrey (Autocorrelation)
    * Whites Test (Heteroskedasticity)
    * AIC - no filter on this 
    * Shapiro-Wilk (Normality of Residuals)
    * Durbin Watson (Autocorrelation) - no filter on this
    inputs: 
    * base: dependent variable and dummies 
    * regs: df with independent variables to be used
    outputs:
    *pass_tests - list of variable names in regs that passed the tests
    * saves file Statistical Test Results.xlsx
    
    next steps: create function for each statistical test and parameterize significance level that is hard-coded here
    '''
    candidates = []
    bgs = []
    whits = []
    sws = []
    pass_tests = []
    aics = []
    dw0 = []
    dw1 = []
    dw2 = []
    dw3 = []
    regs = regs.dropna(axis=1)
    # iterating over the possible independent variables
    for i in regs.columns:
        dep = base['Dependent']
        X = create_basic_design(cp.deepcopy(base), regs, i)
        dep = dep[X.index]
        model = mc.regress_data(dep, X, intercept=True)
        dep_order = residual_integration_order(model.resid, alpha=adf_alpha)
        # creating regression results and diagnostics for each variable
        dep = base['Dependent']
        X = create_design(cp.deepcopy(base), regs, i)
        dep = dep[X.index]
        model = mc.regress_data(dep, X, intercept=True)
        model_image_save(model, i)
        # Serial correlation test
        bg = sm.stats.diagnostic.acorr_breusch_godfrey(model, nlags=4, store=False)  # Null: no autocorrelation
        # Heteroscedasticity Test
        whit = sm.stats.diagnostic.het_white(model.resid, model.model.exog, retres=False)
        # normality test
        sw = shapiro(model.resid)  # Null: Residuals are normally distributed
        # AIC goodness of fit
        aic = model.aic
        # Plot the PACF
        plot_pacf(model.resid, lags=20)
        plt.savefig(i + ' PACF.png')
        plt.close()
        # Save Durbin-Watson pvalues up to 4 lags
        dw = durbin_watson(model, 4)

        # saving regression results and diagnostics for each with significant parameters and I(0) residuals    
        if (dep_order == 0) & ((model.pvalues[len(base.columns):] < param_alpha).all() == True):
            candidates.append(i)
            bgs.append(bg[1])
            whits.append(whit[1])
            sws.append(sw[1])
            aics.append(aic)
            dw0.append(dw[0])
            dw1.append(dw[1])
            dw2.append(dw[2])
            dw3.append(dw[3])
            # filtering for candidates that pass all statistical requirements and plotting 1Q backtest
            if bg[1] > bg_alpha:
                if whit[1] > white_alpha:
                    if sw[1] > sw_alpha:
                        plt.plot(X.index, dep, X.index, model.predict())
                        plt.show()
                        plt.close()
                        pass_tests.append(i)
    results = pd.DataFrame(
        {'Variable': candidates, 'Breusch-Godfrey p': bgs, 'White p': whits, 'Shapiro-Wilk p': sws, 'AIC': aics,
         'DW Lag1': dw0, 'DW Lag2': dw1, 'DW Lag3': dw2, 'DW Lag4': dw3})
    # outputting results to file
    results.to_excel('Statistical Testing Results.xlsx')
    return pass_tests


def create_backtest(X, beta, pq1, figname, target, l=9):
    '''
    create backtesting results and saves .png file for log-log model 
    inputs:
    * X - design matrix for regression
    * beta - parameters from regression
    * pq1 - first forecast period
    * figname - name of figure for file name
    * target - dependent variable target
    outputs:
    * MAPE - MAPE of dynamic backtest plotted
    * saves plot to pwd as [figname].png
    '''
    plt.clf()
    pqloc = target.index.get_loc(pq1)
    hist_X = X.loc[:pq1]
    post_X = X.loc[pq1:]
    post_X = post_X.iloc[0:l]
    y1 = post_X['Lag_Y'].iloc[0]
    # X_hat = post_X.drop('Lag_Y', axis = 1)
    inds = beta.index
    inds = beta.index.drop('Lag_Y')
    inds = inds.drop('const')
    y_hat_values = []
    for j in range(l):
        y_hat = beta['const']
        for i in inds:
            y_hat = y_hat + beta[i] * post_X[i].iloc[j]
        if j == 0:
            y_hat = y_hat + beta['Lag_Y'] * y1
        else:
            y_hat = y_hat + beta['Lag_Y'] * y_hat_values[-1]
        y_hat_values.append(y_hat)
    y_hat_values = pd.Series(np.exp(y_hat_values))

    actual = np.exp(target.iloc[pqloc:pqloc + l])

    y_hat_values.index = actual.index

    error = actual - y_hat_values
    perc_err = error / actual
    ape = abs(perc_err)
    MAPE = np.mean(ape)
    figname = figname + ' MAPE = ' + str(MAPE * 100)[:5] + '%'
    act = pd.Series(np.exp(target.iloc[pqloc - 1]))
    act.index = pd.Series(target.index[pqloc - 1])

    y_hat_values = pd.concat([act, y_hat_values])
    actual = np.exp(target.iloc[0:pqloc + l])

    fig, ax = plt.subplots(ncols=1, nrows=1, figsize=(10, 6))
    plt.ylim(top=1.25 * np.max(y_hat_values), bottom=0.0)
    ax.plot(y_hat_values.index, y_hat_values, actual.index, actual)
    plt.setp(ax.get_xticklabels(), rotation=45)
    ax.legend(['Predicted', 'Actual'], loc='best')
    ax.set_title(figname)

    vals = ax.get_yticks()
    ax.set_yticklabels(['${0:,.0f}'.format(x) for x in vals])
    plt.savefig(figname + '.png')
    plt.close()

    return MAPE


def backtesting(candidates, base, regs, dates, long_dates, target="", l=9):
    '''
    plot and compile backtesting results for candidate models
    inputs:
    * candidates - list of variables to backtest models for
    * base: dependent variable and dummies 
    * regs: df with independent variables to be used
    * dates - list of pq1 periods to start 9Q backtest
    * long_dates - list of pq1 periods to start full length backtest
    * target - target dependent variable
    outputs:
    * mape - matrix of variables and MAPE values for various forecasting periods
    * saves mape matrix to candidate backtesting results.xlsx
    * saves matrix of variables and MAPE values for full backtest to long backtsting results.xlsx file
    '''
    dep = base['Dependent']
    mapes = []
    mape = pd.DataFrame(index=candidates, columns=dates)
    print("MAPE Columns : ", mape.columns)

    long_mape = pd.DataFrame(index=candidates, columns=long_dates)

    for i in candidates:
        X = create_design(base, regs, i)
        dep = dep[X.index]
        model = mc.regress_data(dep, X, intercept=True)
        beta = model.params

        for pq1 in dates:
            figname = i + ' 9Q Backtest ' + str(pq1)
            mape.loc[i, pq1] = create_backtest(X, beta, pq1, figname, base['Dependent'])

        for pq1 in long_dates:
            l = len(X.loc[pq1:].index)
            figname = i + ' Full History Backtest ' + str(pq1)
            long_mape.loc[i, pq1] = create_backtest(X, beta, pq1, figname, base['Dependent'], l=l)
    mape.to_excel('Candidate Backtesting MAPE Results.xlsx')
    long_mape.to_excel('Candidate Long Backtesting MAPE Results.xlsx')

    return mape


def recursive_reg(dep, X, n, varname):
    '''
    perform recursive regression on model 
    inputs:
    * dep - dependent variable
    * X - design matrix
    * n - number of latest periods to use for recursive regression
    * varname - name of the x variable in the model
    outputs:
    * params - parameter values for the recursive regresison trials
    * ps - pvalues for the recursive regression trials
    next steps - change funciton so that a date can be passed in lieu of n
    '''
    # each iteration will generate a set of pvalues and params, then plot each of the pvalues and params.
    ps = pd.DataFrame()
    params = pd.DataFrame()
    for i in range(n, len(dep)):
        dep_trim = dep[0:i]
        X_trim = X.iloc[0:i, :]
        model = mc.regress_data(dep_trim, X_trim, intercept=True)
        params[X_trim.index[-1]] = model.params
        ps[X_trim.index[-1]] = model.pvalues

        for i in params.index:
            if 'Lag_Y' in i:
                plt.ylim(bottom=0.75, top=1.25)
            ax = params.transpose()[i].plot()
            # print("AX", ax)
            ax.set_title(varname + ' model ' + i + ' parameter')
            ax.figure.savefig(varname + ' model ' + i + ' param.png')
            plt.clf()
        for i in ps.index:
            ax = ps.transpose()[i].plot()
            # print("AX", ax)
            ax.set_title(varname + ' model ' + i + ' pvalues')
            ax.figure.savefig(varname + ' model ' + i + ' pval.png')
            plt.clf()

    return params, ps


def create_scenario(mev_full):
    '''
    create mev forecasting scenario
    inputs:
    * mev_full - df of mev data with date as first column
    outputs
    * mev_full - df with date index and no MEV_ prefix on varnames
    '''
    mev_full.index = mev_full.iloc[:, 0]
    mev_full = mev_full.drop(mev_full.columns[0], axis=1)
    for i in mev_full.columns:
        if 'MEV' not in i:
            mev_full = mev_full.drop(i, axis=1)
    cols = []
    for i in mev_full.columns:
        cols.append(i[4:])
    mev_full.columns = cols
    return mev_full


def create_forecast(X, beta, pq1, target_spot, l=22):
    '''
    create the forecast for stress test 
    inputs:
    * X - design matrix for the model
    * beta - parameters of the model
    * pq1 - first quarter to be forecasted
    * target spot - spot value of the dependent variable
    * l - length of forecast desired
    outputs:
    * y_hat_values - predicted outcome of stress test 
    '''
    X = X.iloc[0:l]
    y1 = target_spot
    y_hat = beta[0] + beta[-2] * X.iloc[:, 0] + beta[-1] * X.iloc[:, 1]
    for j in range(l):
        if j == 0:
            y_hat[j] = y_hat[j] + beta[-3] * y1
        else:
            y_hat[j] = y_hat[j] + beta[-3] * y_hat[j - 1]
    y_hat_values = np.exp(y_hat)
    return y_hat_values


def build_scenario(filename, shtm, shtb, shta, shts, beta, pq0, pq1, i, l=22):
    '''
    build forecast for scenarios in the stress test
    inputs:
    * filename - name of the .xlsx file where the data is located
    * shtm - name of the tab with the model development data
    * shtb - name of the tab with the base scenario forecast
    * shta - name of the tab with the adverse scenario forecast
    * shts - name of the tab with the severe scenario forecast
    * beta - parameter vector for the model 
    * pq0 - spot period
    * pq1 - first forecasted period
    * i - model variable name
    * l - length of forecast
    outputs:
    * forecast df - data frame with last 4 periods actuals and the base, adverse, and severe scenario forecasts
    '''
    df = wrangle_model_data(filename, sheet_name=shtm)
    full_dep = df.iloc[:, 0]
    full_dep.name = 'Dependent'
    forecast_df = pd.DataFrame()
    base_full = pd.read_excel(filename, sheet_name=shtb)
    b = create_scenario(base_full)
    adverse_full = pd.read_excel(filename, sheet_name=shta)
    a = create_scenario(adverse_full)
    severe_full = pd.read_excel(filename, sheet_name=shts)
    s = create_scenario(severe_full)
    target_spot = np.log(full_dep[pq0])
    count = 0
    for scen in [b, a, s]:
        count += 1
        X = pd.DataFrame()
        X[i] = scen[i]
        for j in X.index:
            if j not in scen.index:
                X = X.drop(j)
        X.index = scen.index

        lagname = i + ' 1Q Lag'

        lag = scen[i].shift(1)
        X[lagname] = lag
        X = X.loc[pq1:]
        X = np.log(X)

        result = create_forecast(X, beta, pq1, target_spot, l=l)
        forecast_df[str(count)] = result
    forecast_df.columns = ['Base', 'Adverse', 'Severe']
    forecast_df = pd.concat([full_dep, forecast_df], axis=1)
    forecast_df = forecast_df.iloc[-(l + 4):]
    for scenario in ['Base', 'Adverse', 'Severe']:
        forecast_df[scenario].iloc[-(l + 1)] = full_dep.iloc[-1]
    return forecast_df


def stress_test_plot(filename, shtm, shtb, shta, shts, short_list, pq0, pq1, base, regs, bottom="", top=""):
    '''
    plot stress testing results for candidate models
    inputs
    * filename - name of the .xlsx file where the data is located
    * shtm - name of the tab with the model development data
    * shtb - name of the tab with the base scenario forecast
    * shta - name of the tab with the adverse scenario forecast
    * shts - name of the tab with the severe scenario forecast
    * short_list - list of candidate variables for forecasting
    * pq0 - spot period
    * pq1 - first forecasted period
    * base: dependent variable and dummies 
    * regs: df with independent variables to be used
    * bottom - bottom of y-axis of desired forecast graphs
    * top - top of y-axis of desired forecast graphs
    outputs:
    * saves stress test forecast to .png files
    '''
    for i in short_list:
        dep = base['Dependent']
        X = create_design(base, regs, i)
        dep = dep[X.index]
        model = mc.regress_data(dep, X, intercept=True)
        beta = model.params

        forecast = build_scenario(filename, shtm, shtb, shta, shts, beta, pq0, pq1, i)
        figname = i + ' Stress Test Forecast'
        fig, ax = plt.subplots(ncols=1, nrows=1, figsize=(10, 6))

        if top != "":
            if bottom != "":
                plt.ylim(top=top, bottom=bottom)
            else:
                plt.ylim(top=top)
        else:
            if bottom != "":
                plt.ylim(bottom=bottom)

        plt.setp(ax.get_xticklabels(), rotation=45)
        j = 0
        colors = ['Black', 'Green', 'Blue', 'Red']
        for i in forecast.columns:
            ax.plot(forecast.index, forecast[i], color=colors[j])
            j += 1
        #         ax.legend(loc = 'best')
        ax.set_title(figname)
        vals = ax.get_yticks()
        ax.set_yticklabels(['${0:,.0f}'.format(x) for x in vals])
        #         for x in vals:
        #             ax.axhline(y=x, color = 'black', linewidth = 0.2)
        plt.savefig(figname + '.png')
        plt.close()

def wrangle_forecast_data(file, sheet_name):
    ''' 
    create the mev forecasting data frame 
    inputs:
    * file - name of the .xlsx file where the data is located
    * sheet_name - name of the tab with the forecast data scenario
    outputs:
    * df - mev data frame with date index and MEV_ prefix dropped
    '''
    # read the dataset and make date the index of the df
    df = pd.read_excel(file, sheet_name=sheet_name)
    df.index = df['Date']
    df = df.drop('Date', axis=1)
    # Rename columns and remove MEV prefix
    cols = []
    for i in df.columns:
        if 'MEV' in i:
            cols.append(i[4:])
        else:
            cols.append(i)
    df.columns = cols

    return df

def stress_test_compare(filename, shtm, shtb, shta, shts, shtc, short_list, pq0, pq1, base, regs, dep, bottom="",
                        top=""):
    '''
    compare stress test forecasts with those of the current model
    inputs:
    * filename - name of the .xlsx file where the data is located
    * shtm - name of the tab with the model development data
    * shtb - name of the tab with the base scenario forecast
    * shta - name of the tab with the adverse scenario forecast
    * shts - name of the tab with the severe scenario forecast
    * shtc - name of the tab with the comparison forecasts from the current model
    * short_list - list of candidate variables for forecasting
    * pq0 - spot period
    * pq1 - first forecasted period
    * base: dependent variable and dummies 
    * regs: df with independent variables to be used
    * dep - dependent variable 
    * bottom - bottom of y-axis of desired forecast graphs
    * top - top of y-axis of desired forecast graphs
    outputs: 
    * forecast_tbl - forecast metrics for candidate models
    * compare_tbl - forecast metrics for current model
    * saves plots of each forecast to .png file in pwd
    '''
    writer = pd.ExcelWriter('Stress Test Forecast.xlsx', engine='xlsxwriter')
    compare = wrangle_forecast_data(filename, shtc)
    full_df = wrangle_forecast_data(filename, shtm)
    full_dep = full_df.iloc[:, 0]
    full_dep.name = 'Dependent'
    target_spot = full_dep[pq0]
    compare = compare
    for i in short_list:
        X = create_design(base, regs, i)
        dep = dep[X.index]
        model = mc.regress_data(dep, X, intercept=True)
        beta = model.params
        forecast = build_scenario(filename, shtm, shtb, shta, shts, beta, pq0, pq1, i)

        figname = i + ' Stress Test Comparison'
        fig, ax = plt.subplots(ncols=1, nrows=1, figsize=(10, 6))
        if top != "":
            if bottom != "":
                plt.ylim(top=top, bottom=bottom)
            else:
                plt.ylim(top=top)
        else:
            if bottom != "":
                plt.ylim(bottom=bottom)

        plt.setp(ax.get_xticklabels(), rotation=45)
        j = 0
        colors = ['Black', 'Green', 'Blue', 'Red']
        forecast.columns = ['Actual', 'Model Base', 'Model_Adverse', 'Model_Severe']
        forecast_tbl = pd.DataFrame()
        for k in forecast.columns:
            if j > 0:
                actual = target_spot
                pq4 = forecast[k].iloc[4]
                pq9 = forecast[k].iloc[9]
                pq20 = forecast[k].iloc[20]
                cagr9 = (pq9 / actual) ** (4 / 9) - 1
                cagr20 = (pq20 / actual) ** (4 / 20) - 1

                actual = '${0:,.0f}'.format(actual)
                pq4 = '${0:,.0f}'.format(pq4)
                pq9 = '${0:,.0f}'.format(pq9)
                pq20 = '${0:,.0f}'.format(pq20)
                cagr9 = '{:.2%}'.format(cagr9)
                cagr20 = '{:.2%}'.format(cagr20)
                forecast_tbl[k] = pd.Series([actual, pq4, pq9, pq20, cagr9, cagr20])
            ax.plot(forecast.index, forecast[k], color=colors[j])
            j += 1
        forecast_tbl = forecast_tbl.transpose()
        forecast_tbl.columns = ['Actual', 'PQ4', 'PQ9', 'PQ20', '9Q CAGR', '20Q CAGR']

        for col in compare.columns:
            compare.loc[full_dep.index[-1], col] = target_spot

        compare = compare.sort_index(axis=0)

        colors = ['Green', 'Blue', 'Red']
        j = 0
        compare_tbl = pd.DataFrame()
        for k in compare.columns:
            ax.plot(compare.index, compare[k], color=colors[j], linestyle='--')

            actual = target_spot
            pq4 = compare[k].iloc[4]
            pq9 = compare[k].iloc[9]
            pq20 = compare[k].iloc[20]
            cagr9 = (pq9 / actual) ** (4 / 9) - 1
            cagr20 = (pq20 / actual) ** (4 / 20) - 1

            actual = '${0:,.0f}'.format(actual)
            pq4 = '${0:,.0f}'.format(pq4)
            pq9 = '${0:,.0f}'.format(pq9)
            pq20 = '${0:,.0f}'.format(pq20)
            cagr9 = '{:.2%}'.format(cagr9)
            cagr20 = '{:.2%}'.format(cagr20)
            compare_tbl[k] = pd.Series([actual, pq4, pq9, pq20, cagr9, cagr20])
            j += 1
        compare_tbl = compare_tbl.transpose()
        compare_tbl.columns = ['Actual', 'PQ4', 'PQ9', 'PQ20', '9Q CAGR', '20Q CAGR']
        #         ax.legend(loc = 'best')
        ax.set_title(figname)
        vals = ax.get_yticks()
        ax.set_yticklabels(['${0:,.0f}'.format(x) for x in vals])
        #         for x in vals:
        #             ax.axhline(y=x, color = 'black', linewidth = 0.2)
        plt.savefig(figname + '.png')
        plt.close()

        forecast_tbl.to_excel(writer, sheet_name=i[:25] + " New")
        compare_tbl.to_excel(writer, sheet_name=i[:25] + ' Old')
    writer.save()
    return forecast_tbl, compare_tbl


def create_rands(size, n):
    '''
    generate vector 1xsize random numbers between -n and n 
    inputs:
    * size - size of vector desired
    * n - number for random numbers between -n and n 
    outputs:
    * rands - series of randomly generated numbers
    '''
    rands = []
    for i in range(size):
        x = random.randrange(-n * 1000, n * 1000, 1) / 1000
        rands.append(x)
    return pd.Series(rands)


def create_random_forecast(X, beta, y1, xspot, xmu, xsig, l=12):
    '''
    create random forecasts based on random variance added to mev forecast
    x_t = x_{t-1} + historical mean(x') +  random decimal (-3,3) * historical standard deviation(x')
    where ' denotes a difference
    inputs:
    * X - design matrix of the model
    * beta - parameters of the model
    * xspot - spot value of the regressor variable x
    * xmu - historical mean of the diff of the x variable 
    * xsig - historical standard deviation of the diff of the x variable
    l - length of forecast
    outputs: 
    * y_hat_values - model forecast for random scenario
    '''
    X = X.iloc[0:l]
    xsigrands = create_rands(l, 3)
    xlagsigrands = create_rands(l, 3)
    xsigrands.index = X.index
    xlagsigrands.index = X.index
    x = []
    xlag = []
    for i in range(l):
        if i == 0:
            x.append(xspot + xmu + xsigrands[i] * xsig)  # + xmu)
            xlag.append(xspot)
        else:
            x.append(x[-1] + xmu + xsigrands[i] * xsig)  # + xmu)
            xlag.append(x[-1])

    y_hat = beta[0] + beta[-2] * pd.Series(x) + beta[-1] * pd.Series(xlag)
    for j in range(l):
        if j == 0:
            y_hat[j] = y_hat[j] + beta[-3] * y1
        else:
            y_hat[j] = y_hat[j] + beta[-3] * y_hat[j - 1]

    y_hat_values = np.exp(y_hat.astype(float))
    y_hat_values = pd.Series(y_hat_values)
    y_hat_values.index = X.index
    return y_hat_values


def build_random_scenario(filename, shtm, shtb, beta, pq0, pq1, i, l=12):
    '''
    build random forecast for the model
    inputs:
    * filename - name of the .xlsx file where the data is located
    * shtm - name of the tab with the model development data
    * shtb - name of the tab with the base scenario forecast
    * beta - parameters of the model
    * pq0 - spot period
    * pq1 - first forecasted period
    * i - variable for the model
    * l - length of forecast period
    outputs:
    * forecast_df - random model forecast
    '''
    forecast_df = pd.DataFrame()
    base_full = pd.read_excel(filename, sheet_name=shtb)
    scen = create_scenario(base_full)
    count = 0
    X = pd.DataFrame()
    X[i] = scen[i]

    #     for j in X.index:
    #         if j not in scen.index:
    #             X = X.drop(j)
    X.index = scen.index
    lagname = i + ' 1Q Lag'
    lag = scen[i].shift(1)
    X[lagname] = lag

    df = wrangle_model_data(filename, shtm)
    dep = df.iloc[:, 0]
    dep.name = 'Dependent'
    y1 = np.log(dep[pq0])
    X = np.log(X)
    diff = X - X.shift(1)
    diff = diff.loc[df.index]
    xsig = np.std(diff[i])
    xmu = np.mean(diff[i])
    xspot = X[i].loc[pq0]
    X = X.loc[pq1:]
    result = create_random_forecast(X, beta, y1, xspot, xmu, xsig)
    forecast_df['Rand'] = result

    forecast_df = pd.concat([dep, forecast_df], axis=1)
    forecast_df = forecast_df.iloc[-(l + 4):]
    forecast_df['Rand'].iloc[-(l + 1)] = forecast_df['Dependent'].iloc[-(l + 1)]
    return forecast_df


def create_sensitivity(filename, shtm, shtb, base, regs, short_list, pq0, pq1, l=50):
    '''
    create sensitivity testing for list of candidate models
    inputs: 
    * filename - name of the .xlsx file where the data is located
    * shtm - name of the tab with the model development data
    * shtb - name of the tab with the base scenario forecast
    * base: dependent variable and dummies 
    * regs: df with independent variables to be used
    * short_list - list of candidate variables for forecasting
    * pq0 - spot period
    * pq1 - first forecasted period
    * l - number of forecasts to randomly generate
    outputs:
    * saves plot of sensitivity analysis to .png in pwd
    '''
    resdf = pd.DataFrame(columns=['Mean', 'Min', 'Max'])
    df = wrangle_model_data(filename, shtm)
    dep = df.iloc[:, 0]
    dep.name = 'Dependent'
    dep = np.log(dep)
    for i in short_list:
        X = create_design(base, regs, i)
        dep = dep[X.index]
        model = mc.regress_data(dep, X, intercept=True)
        beta = model.params
        figname = i + ' Dynamic Sensitivity Testing'
        fig, ax = plt.subplots(ncols=1, nrows=1, figsize=(10, 6))
        ends = []
        for k in range(l):
            j = 0
            forecast = build_random_scenario(filename, shtm, shtb, beta, pq0, pq1, i)
            ends.append(forecast['Rand'].iloc[-1])
            colors = ['black', 'red']
            for col in forecast.columns:
                ax.plot(forecast.index, forecast[col], color=colors[j])
                j += 1
        try:
            plt.ylim(top=1.25 * np.max(np.max(forecast)), bottom=0.75 * np.min(np.min(forecast)))
        except:
            pass
        plt.setp(ax.get_xticklabels(), rotation=45)
        ax.set_title(figname)
        vals = ax.get_yticks()
        ax.set_yticklabels(['${0:,.0f}'.format(x) for x in vals])
        for x in vals:
            ax.axhline(y=x, color='black', linewidth=0.2)
        plt.savefig(figname + '.png')
        plt.close()

        resdf.loc[i] = ['${0:,.0f}'.format(np.mean(ends)), '${0:,.0f}'.format(np.min(ends)),
                        '${0:,.0f}'.format(np.max(ends))]
    writer = pd.ExcelWriter('Sensitivity Metrics.xlsx', engine='xlsxwriter')
    resdf.to_excel(writer)
    writer.save()

# Save regression and backtesting results of candidate variables to file
def compile_results(short_list):
    '''
    create a word doc of the key results for the final candidate models
    inputs:
    * short_list - list of candidates to compile to document
    outputs:
    * saves .docx file to pwd with results compiled and organized
    '''
    try:
        os.remove('short_list.docx')
    except:
        pass

    document = Document()
    document.add_heading('Document Title', 0)
    for folder in short_list:
        document.add_heading(folder, level=1)
        ls = os.listdir()
        varfiles = []
        for i in ls:
            if (folder in i) and ('VS. ' + folder not in i):
                varfiles.append(i)
        modelfiles = []
        backfiles = []
        outfiles = []
        stressfiles = []
        sensfiles = []
        paramfiles = []
        pfiles = []
        pacffiles = []
        for file in varfiles:
            if file.endswith(".txt"):
                modelfiles.append(file)
            if 'PACF' in file:
                pacffiles.append(file)
            if 'Backtest' in file:
                backfiles.append(file)
            if 'Out of Time' in file:
                outfiles.append(file)
            if 'Stress Test' in file:
                stressfiles.append(file)
            if 'Sensitivity' in file:
                sensfiles.append(file)
            if 'param.png' in file:
                paramfiles.append(file)
            if 'pval.png' in file:
                pfiles.append(file)
        document.add_heading('Model Form', level=2)
        for file in modelfiles:
            file1 = open(file, "r")
            graph = file1.read()
            p = document.add_paragraph(graph)
        for file in pacffiles:
            document.add_picture(file, width=Inches(7))
        document.add_heading('Backtesting', level=2)
        for file in backfiles:
            document.add_picture(file, width=Inches(7))
        document.add_heading('Out of Time Testing', level=2)
        for file in outfiles:
            document.add_picture(file, width=Inches(7))
        document.add_heading('Stress Testing', level=2)
        for file in stressfiles:
            document.add_picture(file, width=Inches(7))
        document.add_heading('Sensitivity Testing', level=2)
        for file in sensfiles:
            document.add_picture(file, width=Inches(7))
        document.add_heading('Recursive', level=2)
        document.add_heading('Parameter Plots', level=3)
        for file in paramfiles:
            document.add_picture(file, width=Inches(7))
        document.add_heading('P-Value Plots', level=3)
        for file in pfiles:
            document.add_picture(file, width=Inches(7))
        document.add_page_break()
    document.save('short_list.docx')


def copy_output(folder):
    '''
    copy docx and xlsx files to folder
    '''
    try:
        os.mkdir(folder)
    except:
        pass
    for i in os.listdir():
        if i.endswith('.xlsx'):
            shutil.copy(i, folder)
        if i.endswith('.docx'):
            shutil.copy(i, folder)


def version_output(folder):
    ''' 
    organize output into folders
    '''
    os.mkdir(folder)
    modelloc = folder + '/model_forms'
    os.mkdir(modelloc)
    backtestloc = folder + '/backtests'
    os.mkdir(backtestloc)
    stloc = folder + '/stress_test'
    os.mkdir(stloc)
    sensloc = folder + '/sensitivity_analysis'
    os.mkdir(sensloc)
    paramloc = folder + '/recursive_parameters'
    os.mkdir(paramloc)
    ploc = folder + '/recursive_ pvalues'
    os.mkdir(ploc)
    statsloc = folder + '/statistical_testing'
    os.mkdir(statsloc)
    pacfloc = folder + '/pacf'
    os.mkdir(pacfloc)
    ootloc = folder + '/out_of_time'
    os.mkdir(ootloc)

    for i in os.listdir():
        if i.endswith('.txt'):
            shutil.move(i, modelloc)
        if 'Out of Time' in i:
            shutil.move(i, ootloc)
        if 'Backtest' in i:
            shutil.move(i, backtestloc)
        if 'Stress Test' in i:
            shutil.move(i, stloc)
        if 'Sensitivity' in i:
            shutil.move(i, sensloc)
        if 'param' in i:
            shutil.move(i, paramloc)
        if 'pval' in i:
            shutil.move(i, ploc)
        if 'PACF' in i:
            shutil.move(i, pacfloc)
        if 'Statistical' in i:
            shutil.move(i, statsloc)
        if i.endswith('.docx'):
            try:
                shutil.move(i, folder)
            except:
                pass


def oot_backtesting(candidates, base, regs, pq0, dates, full_base, full_regs, target, l=9):
    '''
    perform out of time testing on the list of candidate models
    inputs:
    * candidates - list of models to perform oot testing on
    * base: dependent variable and dummies for intended development period for oot test
    * regs: df with independent variables for intended development period for oot test
    * pq0 - spot period
    * dates - dates to perform backtesting on 
    * full base - full dependent variables and dummies 
    * full_regs - full independent variables
    * target - model target actuals
    * l - length of forecast for backtesting
    outputs:
    * mape - matrix of MAPE values for forecasts indexed by pq1 period
    '''
    dep = base['Dependent']
    mapes = []
    mape = pd.DataFrame(index=candidates, columns=dates)
    for i in candidates:
        full_X = create_design(full_base, full_regs, i)
        X = create_design(base, regs, i)
        dep = dep[X.index]
        model = mc.regress_data(dep, X, intercept=True)
        beta = model.params
        for pq1 in dates:
            figname = i + '_' + str(pq0)[:11] + ' Out of Time ' + str(pq1)
            mape.loc[i, pq1] = create_backtest(full_X, beta, pq1, figname, target)
    mape.to_excel('Out of Time MAPE Results.xlsx')
    return mape


def out_of_time(candidates, base, regs, pq0, dates):
    '''
    create out of time testing for list of candidate variables
    inputs: 
    * candidates - list of models to perform oot testing on
    * base: dependent variable and dummies 
    * regs: df with independent variables 
    * pq0 - spot period to use for out of time regressions
    * dates - pq1 periods to perform backtesting on
    outputs:
    * mape - matrix of MAPE values for forecasts indexed by pq1 period
    '''
    full_dep = base['Dependent']
    full_base = base
    full_regs = regs
    base = base[:pq0]
    regs = regs[:pq0]
    mape = oot_backtesting(candidates, base, regs, pq0, dates, full_base, full_regs, full_dep, l=9)
    return mape


#######################################
# ABC "interface" classes             #
#######################################

#######################################
# Abstract classes                    #
#######################################

#######################################
# Concrete classes                    #
#######################################

#######################################
# Initialization needed after member  #
#   definition is complete            #
#######################################

#######################################
# Imports needed after member         #
#   definition (to resolve circular   #
#   dependencies - avoid if at all    #
#   possible                          #
#######################################

#######################################
# Code to execute if the module is    #
#   called directly                   #
#######################################

if __name__ == '__main__':
    pass
